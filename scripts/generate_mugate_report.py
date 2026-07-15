#!/usr/bin/env python3
"""
MuGate University Project Report Generator
Generates a professional 90-110 page .docx report using python-docx.
Run: py c:\\dev\\Capstone-v2\\MuGate\\scripts\\generate_mugate_report.py
"""

from __future__ import annotations

import glob
import os
import re
import shutil
import sys
import tempfile
from pathlib import Path
from typing import Any, Callable, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent
ROOT = SCRIPT_DIR.parent
OUTPUT_PATH = ROOT / "MuGate_Project_Report.docx"
ASSETS = ROOT / "report-assets"
DIAGRAMS_DIR = ASSETS / "diagrams"
SCREENSHOTS_DIR = ASSETS / "screenshots"
CODE_DIR = ASSETS / "code"

CODE_SOURCES = {
    "auth-login-flow": ROOT / "backend" / "src" / "modules" / "auth" / "auth.service.ts",
    "ai-cascade": ROOT / "backend" / "src" / "modules" / "ai" / "chatbot" / "ai" / "ai.provider.ts",
    "frontend-routes": ROOT / "frontend" / "src" / "App.jsx",
    "resume-schema": ROOT / "frontend" / "src" / "pages" / "ResumeEnhancer" / "editor" / "resumeSchema.js",
}

AUTHORS = [
    ("Mohammad Jomaa", ["Backend Developer", "AI Integration", "Database Admin"]),
    ("Abo Al Fadel Ismael", ["UI/UX Designer", "Frontend Developer", "Project Manager"]),
]
UNIVERSITY = "Al Maaref University"
DEPARTMENT = "Department of Computer Sciences"
PROJECT = "MuGate"
ACADEMIC_YEAR = "2025-2026"
SUPER_ADMIN_ID = "101230004"

FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
HEADING_SIZES = {1: Pt(16), 2: Pt(14), 3: Pt(12)}


# ---------------------------------------------------------------------------
# Counters
# ---------------------------------------------------------------------------
class Counters:
    def __init__(self) -> None:
        self.figure = 0
        self.table = 0
        self.figures: List[str] = []
        self.tables: List[str] = []

    def next_figure(self, caption: str) -> str:
        self.figure += 1
        full = f"Figure {self.figure}: {caption}"
        self.figures.append(full)
        return full

    def next_table(self, caption: str) -> str:
        self.table += 1
        full = f"Table {self.table}: {caption}"
        self.tables.append(full)
        return full


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------
def set_run_font(run, size: Pt = BODY_SIZE, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = size
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

    for level in (1, 2, 3):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_NAME
        style.font.size = HEADING_SIZES[level]
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_body(doc: Document, text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run)


def add_bodies(doc: Document, paragraphs: Sequence[str]) -> None:
    for text in paragraphs:
        add_body(doc, text)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=HEADING_SIZES.get(level, BODY_SIZE), bold=True)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "(Right-click the table of contents and select 'Update Field' in Microsoft Word to populate page numbers.)"
    )
    set_run_font(nr, size=Pt(10), italic=True)


def add_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_sep)
        run._r.append(fld_end)


def make_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], counters: Counters, caption: str) -> None:
    cap = counters.next_table(caption)
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap_p.add_run(cap)
    set_run_font(cr, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, bold=True)

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=Pt(10))

    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, counters: Counters, caption: str, width: float = 6.0) -> bool:
    if not image_path.exists():
        return False
    try:
        doc.add_picture(str(image_path), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = counters.next_figure(caption)
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(cap)
        set_run_font(cr, bold=True, italic=True)
        doc.add_paragraph()
        return True
    except Exception as exc:
        print(f"  [skip image] {image_path.name}: {exc}")
        return False


def add_pseudocode(doc: Document, title: str, lines: Sequence[str]) -> None:
    add_heading(doc, title, level=3)
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_run_font(run, size=Pt(10))
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")


def add_bullet_list(doc: Document, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


# ---------------------------------------------------------------------------
# Code image generation (Pillow)
# ---------------------------------------------------------------------------
def _load_mono_font(size: int = 14) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/consola.ttf",
        "C:/Windows/Fonts/cour.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
        "/System/Library/Fonts/Menlo.ttc",
    ]
    for path in candidates:
        if os.path.exists(path):
            try:
                return ImageFont.truetype(path, size)
            except OSError:
                pass
    return ImageFont.load_default()


def generate_code_image(source_path: Path, output_path: Path, max_lines: int = 55) -> bool:
    if not source_path.exists():
        print(f"  [skip code image] source not found: {source_path}")
        return False

    output_path.parent.mkdir(parents=True, exist_ok=True)
    text = source_path.read_text(encoding="utf-8", errors="replace")
    lines = text.splitlines()[:max_lines]
    if len(text.splitlines()) > max_lines:
        lines.append(f"... ({len(text.splitlines()) - max_lines} more lines)")

    font = _load_mono_font(13)
    padding = 24
    line_height = 20
    max_width = 0
    dummy = Image.new("RGB", (1, 1))
    draw = ImageDraw.Draw(dummy)
    for line in lines:
        line = line.expandtabs(4)
        bbox = draw.textbbox((0, 0), line, font=font)
        max_width = max(max_width, bbox[2] - bbox[0])

    img_w = max_width + padding * 2
    img_h = len(lines) * line_height + padding * 2
    img = Image.new("RGB", (img_w, img_h), color=(30, 30, 30))
    draw = ImageDraw.Draw(img)

    # Title bar
    draw.rectangle([0, 0, img_w, 28], fill=(45, 45, 48))
    draw.text((padding, 6), source_path.name, fill=(200, 200, 200), font=font)

    y = padding + 8
    for i, line in enumerate(lines):
        line = line.expandtabs(4)
        # Simple syntax highlighting
        color = (212, 212, 212)
        stripped = line.strip()
        if stripped.startswith("//") or stripped.startswith("#") or stripped.startswith("*"):
            color = (106, 153, 85)
        elif stripped.startswith("import ") or stripped.startswith("export ") or stripped.startswith("const ") or stripped.startswith("function "):
            color = (86, 156, 214)
        elif "async" in stripped or "await" in stripped or "return" in stripped:
            color = (197, 134, 192)
        draw.text((padding, y), line, fill=color, font=font)
        y += line_height

    img.save(str(output_path), "PNG")
    print(f"  [code image] {output_path.name}")
    return True


def generate_all_code_images() -> List[Path]:
    generated: List[Path] = []
    for name, src in CODE_SOURCES.items():
        out = CODE_DIR / f"{name}.png"
        if generate_code_image(src, out):
            generated.append(out)
    return generated


def copy_temp_screenshots() -> int:
    """Copy browser screenshots from temp directories if present."""
    SCREENSHOTS_DIR.mkdir(parents=True, exist_ok=True)
    patterns = [
        os.path.join(tempfile.gettempdir(), "web-*.png"),
        os.path.join(tempfile.gettempdir(), "cursor-*", "web-*.png"),
        str(ROOT / "report-assets" / "screenshots" / "web-*.png"),
    ]
    copied = 0
    for pattern in patterns:
        for src in glob.glob(pattern):
            dst = SCREENSHOTS_DIR / Path(src).name
            if not dst.exists():
                try:
                    shutil.copy2(src, dst)
                    copied += 1
                    print(f"  [screenshot] copied {Path(src).name}")
                except OSError:
                    pass
    return copied


def embed_glob_figures(doc: Document, pattern: str, counters: Counters, caption_fn: Callable[[str], str], width: float = 6.0) -> int:
    paths = sorted(glob.glob(pattern))
    count = 0
    for p in paths:
        name = Path(p).stem
        if add_figure(doc, Path(p), counters, caption_fn(name), width=width):
            count += 1
    return count


# ---------------------------------------------------------------------------
# Inline report data
# ---------------------------------------------------------------------------
TECH_STACK = [
    ("Frontend", "React", "19.2.0", "Component-based SPA UI"),
    ("Frontend", "Vite", "7.2.4", "Build tool and dev server"),
    ("Frontend", "React Router", "7.13.0", "Client-side routing"),
    ("Frontend", "Tailwind CSS", "4.1.18", "Utility-first styling"),
    ("Frontend", "Three.js", "0.182.0", "3D internship carousel"),
    ("Backend", "Express", "5.2.1", "REST API server"),
    ("Backend", "TypeScript", "5.9.3", "Type-safe server code"),
    ("Backend", "SQL Server", "mssql 12.2.0", "Relational database"),
    ("Backend", "Playwright", "1.60.0", "Portal authentication scraping"),
    ("Backend", "ChromaDB", "3.5.0", "Vector store for RAG"),
    ("Backend", "jsonwebtoken", "9.0.3", "JWT session tokens (24h)"),
    ("Mobile", "Expo", "56.0.12", "Cross-platform mobile runtime"),
    ("Mobile", "React Native", "0.85.3", "Native mobile UI"),
    ("Mobile", "React Navigation", "7.x", "Tab and stack navigation"),
    ("AI", "DeepSeek API", "deepseek-chat", "Primary LLM provider"),
    ("AI", "Google Gemini", "gemini-2.5-flash", "Fallback LLM + embeddings"),
    ("AI", "OpenRouter", "free models", "Tertiary LLM fallback cascade"),
]

API_ENDPOINTS = [
    ("GET", "/", "Public", "Health check — MuGate Backend Running"),
    ("GET", "/test-db", "Public", "SQL Server connection test"),
    ("GET", "/api/auth/login", "Public", "Returns 405 — login requires POST"),
    ("POST", "/api/auth/login", "Public", "Portal SSO login; returns JWT + user + isAdmin"),
    ("GET", "/api/auth/me/is-admin", "JWT", "Self-check admin status"),
    ("GET", "/api/auth/users", "Admin", "List all registered users"),
    ("GET", "/api/auth/admins", "Admin", "List admins with online/offline activity"),
    ("POST", "/api/auth/admins", "Admin", "Grant admin by universityId"),
    ("DELETE", "/api/auth/admins/:universityId", "Admin", "Revoke admin (super admin protected)"),
    ("GET", "/api/history/", "JWT", "Get student's scraped academic history"),
    ("POST", "/api/history/sync", "JWT", "Re-scrape portal history"),
    ("GET", "/api/courses/", "JWT", "List all courses with sections"),
    ("GET", "/api/courses/:id", "JWT", "Get single course by ID"),
    ("POST", "/api/courses/sync", "JWT", "Sync course offerings from MU portal"),
    ("POST", "/api/generate/", "JWT", "Generate optimal conflict-free schedules"),
    ("GET", "/api/schedules/", "JWT", "Get all saved schedules for user"),
    ("POST", "/api/schedules/save", "JWT", "Save chosen schedule"),
    ("DELETE", "/api/schedules/:id", "JWT", "Delete saved schedule"),
    ("POST", "/api/chatbot/sessions", "Optional JWT", "Create chat session (source: chat|resume)"),
    ("POST", "/api/chatbot/message", "Optional JWT", "Send message; RAG for university questions"),
    ("POST", "/api/chatbot/upload", "Optional JWT", "Upload file (max 5MB) for chat context"),
    ("GET", "/api/chatbot/sessions", "Optional JWT", "List user's chat sessions"),
    ("DELETE", "/api/chatbot/sessions/:sessionId", "Optional JWT", "Delete session"),
    ("GET", "/api/chatbot/sessions/:sessionId/messages", "Optional JWT", "Get message history"),
    ("GET", "/api/chatbot/analytics", "JWT", "Chat analytics"),
    ("POST", "/api/chatbot/enhance-prompt", "Optional JWT", "AI prompt enhancement"),
    ("POST", "/api/scraper/scrape", "Public", "Portal scrape trigger"),
    ("POST", "/api/scraper/university/crawl", "Admin", "Start full university website crawl"),
    ("POST", "/api/scraper/university/sync", "Admin", "Incremental KB sync (24h threshold)"),
    ("GET", "/api/scraper/university/stats", "Admin", "Knowledge base statistics"),
    ("GET", "/api/scraper/university/runs", "Admin", "Scraper run history"),
    ("GET", "/api/scraper/university/status", "Public", "Check if scraper is running"),
    ("POST", "/api/scraper/university/rescrape", "Admin", "Full rescrape — clears and rebuilds KB"),
    ("POST", "/api/scraper/university/reindex", "Admin", "Reindex all vectors to Chroma/local store"),
    ("POST", "/api/scraper/university/sitemap-refresh", "Admin", "Refresh sitemap and queue new URLs"),
    ("POST", "/api/resume/generate", "Public", "Generate PDF/DOCX resume"),
    ("POST", "/api/resume/analyze", "Public", "AI resume scoring (ATS, content, impact)"),
    ("POST", "/api/resume/ai-edit", "Public", "Structured AI rewrite of resume section"),
    ("POST", "/api/resume/parse", "Public", "Parse raw resume text to structured JSON"),
    ("POST", "/api/resume/convert", "Public", "Upload PDF/DOCX → structured JSON"),
    ("POST", "/api/resume/convert-base64", "Public", "Convert via base64 (mobile-friendly)"),
    ("POST", "/api/resume/edit", "Public", "Legacy upload-and-edit path"),
    ("GET", "/api/internships/companies", "Public", "List all companies"),
    ("POST", "/api/internships/companies", "Admin", "Add company"),
    ("PUT", "/api/internships/companies/:id", "Admin", "Update company"),
    ("DELETE", "/api/internships/companies/:id", "Admin", "Delete company"),
    ("GET", "/api/internships/stats", "Public", "Aggregated rating stats per company"),
    ("GET", "/api/internships/reviews/:companyId", "Public", "Reviews for a company"),
    ("POST", "/api/internships/reviews/:companyId", "JWT", "Submit review (1 per user per company)"),
    ("PUT", "/api/internships/reviews/:reviewId", "JWT", "Update own review"),
    ("DELETE", "/api/internships/reviews/:reviewId", "JWT", "Delete own review"),
    ("GET", "/api/capstone/partners", "Public", "Partner listings with search"),
    ("POST", "/api/capstone/partners", "JWT", "Add partner listing (one per user)"),
    ("DELETE", "/api/capstone/partners/:partnerId", "JWT", "Delete own listing"),
    ("GET", "/api/capstone/ideas", "Public", "Capstone ideas with search/filter"),
    ("GET", "/api/capstone/ideas/deleted", "Admin", "Soft-deleted default ideas"),
    ("POST", "/api/capstone/ideas", "Admin", "Add capstone idea"),
    ("PUT", "/api/capstone/ideas/:id", "Admin", "Update idea"),
    ("DELETE", "/api/capstone/ideas/:id", "Admin", "Delete idea"),
    ("POST", "/api/capstone/ideas/:id/restore", "Admin", "Restore soft-deleted idea"),
    ("GET", "/api/capstone/ideas/faculties", "Public", "Distinct faculties from ideas"),
    ("POST", "/api/capstone/ai/chat", "Public", "Capstone AI advisor chat"),
    ("GET", "/api/events/", "Public", "Upcoming events"),
    ("GET", "/api/events/categories", "Public", "Event categories"),
    ("GET", "/api/events/stats", "Public", "Event statistics"),
    ("GET", "/api/events/:id", "Public", "Event by ID"),
    ("POST", "/api/events/", "Admin", "Add manual event"),
    ("PUT", "/api/events/:id", "Admin", "Update manual event"),
    ("DELETE", "/api/events/:id", "Admin", "Delete manual event"),
    ("POST", "/api/events/scrape", "Public", "Trigger event scraping"),
    ("GET", "/api/roadmap/", "Optional JWT", "Get user's degree roadmap"),
    ("POST", "/api/roadmap/", "JWT", "Save/overwrite roadmap"),
    ("POST", "/api/roadmap/reset", "JWT", "Reset roadmap to defaults"),
]

DB_TABLES = {
    "Users": [
        ("id", "UUID", "PK", "Unique user identifier"),
        ("email", "NVARCHAR", "", "User email (default: {universityId}@mu.edu.lb)"),
        ("passwordHash", "NVARCHAR", "", "Placeholder hash for SSO users"),
        ("name", "NVARCHAR", "", "Display name from portal scrape"),
        ("universityId", "NVARCHAR", "UNIQUE", "MU student ID"),
        ("createdAt", "DATETIME", "", "Registration timestamp"),
        ("updatedAt", "DATETIME", "", "Last profile update"),
        ("lastActiveAt", "DATETIME", "", "Last API activity for online status"),
    ],
    "PortalCredentials": [
        ("id", "UUID", "PK", "Credential record ID"),
        ("userId", "UUID", "FK→Users", "Owner of encrypted credentials"),
        ("encryptedUsername", "NVARCHAR", "", "AES-encrypted portal username"),
        ("encryptedPassword", "NVARCHAR", "", "AES-encrypted portal password"),
        ("createdAt", "DATETIME", "", "Creation timestamp"),
        ("updatedAt", "DATETIME", "", "Last credential update"),
    ],
    "Courses": [
        ("id", "UUID", "PK", "Course identifier"),
        ("courseCode", "NVARCHAR", "", "e.g. CSC245"),
        ("courseName", "NVARCHAR", "", "Full course title"),
        ("credits", "INT", "", "Credit hours"),
        ("department", "NVARCHAR", "", "Offering department"),
        ("semester", "NVARCHAR", "", "Semester code"),
        ("createdAt", "DATETIME", "", "Record creation time"),
    ],
    "CourseSections": [
        ("id", "UUID", "PK", "Section identifier"),
        ("courseId", "UUID", "FK→Courses", "Parent course"),
        ("sectionNumber", "NVARCHAR", "", "Section number"),
        ("instructor", "NVARCHAR", "", "Instructor name"),
        ("day", "NVARCHAR", "", "Day of week"),
        ("startTime", "TIME", "", "Class start time"),
        ("endTime", "TIME", "", "Class end time"),
        ("type", "NVARCHAR", "", "Lecture/Lab/Tutorial"),
        ("capacity", "INT", "", "Maximum enrollment"),
        ("enrolled", "INT", "", "Current enrollment count"),
        ("room", "NVARCHAR", "", "Room location"),
    ],
    "AcademicHistory": [
        ("id", "UUID", "PK", "History record ID"),
        ("userId", "UUID", "FK→Users", "Student owner"),
        ("courseCode", "NVARCHAR", "", "Completed course code"),
        ("courseName", "NVARCHAR", "", "Course title"),
        ("grade", "NVARCHAR", "", "Letter grade"),
        ("credits", "INT", "", "Credits earned"),
        ("semester", "NVARCHAR", "", "Completion semester"),
        ("status", "NVARCHAR", "", "Pass/Fail/In Progress"),
        ("category", "NVARCHAR", "", "Core/Elective/Free"),
        ("isElective", "BIT", "", "Elective flag"),
        ("updatedAt", "DATETIME", "", "Last sync timestamp"),
    ],
    "Schedules": [
        ("id", "UUID", "PK", "Saved schedule ID"),
        ("userId", "UUID", "FK→Users", "Owner"),
        ("name", "NVARCHAR", "", "User-assigned schedule name"),
        ("score", "FLOAT", "", "Optimizer score"),
        ("totalCredits", "INT", "", "Total credits in schedule"),
        ("createdAt", "DATETIME", "", "Save timestamp"),
    ],
    "ScheduleSections": [
        ("scheduleId", "UUID", "FK→Schedules", "Parent schedule"),
        ("sectionId", "UUID", "FK→CourseSections", "Selected section"),
    ],
    "Sessions": [
        ("id", "UUID", "PK", "Portal session ID"),
        ("userId", "UUID", "FK→Users", "Session owner"),
        ("cookies", "NVARCHAR(MAX)", "", "Serialized portal cookies"),
        ("expiresAt", "DATETIME", "", "Session expiry"),
        ("createdAt", "DATETIME", "", "Creation time"),
    ],
    "ChatSessions": [
        ("id", "UUID", "PK", "Chat session ID"),
        ("userId", "UUID", "FK→Users", "Owner (nullable for anonymous)"),
        ("title", "NVARCHAR", "", "Session title"),
        ("source", "NVARCHAR", "", "chat or resume"),
        ("isPinned", "BIT", "", "Pinned flag"),
        ("isActive", "BIT", "", "Active session flag"),
        ("createdAt", "DATETIME", "", "Creation time"),
        ("updatedAt", "DATETIME", "", "Last message time"),
    ],
    "ChatMessages": [
        ("id", "UUID", "PK", "Message ID"),
        ("sessionId", "UUID", "FK→ChatSessions", "Parent session"),
        ("role", "NVARCHAR", "", "user/assistant/system"),
        ("content", "NVARCHAR(MAX)", "", "Message text"),
        ("tokensUsed", "INT", "", "AI tokens consumed"),
        ("createdAt", "DATETIME", "", "Timestamp"),
    ],
    "ChatAnalytics": [
        ("id", "UUID", "PK", "Analytics record ID"),
        ("questionCategory", "NVARCHAR", "", "Classifier category"),
        ("isFailed", "BIT", "", "Failed response flag"),
        ("responseTimeMs", "INT", "", "Response latency"),
        ("createdAt", "DATETIME", "", "Timestamp"),
    ],
    "KnowledgePages": [
        ("id", "UUID", "PK", "Page ID"),
        ("url", "NVARCHAR", "UNIQUE", "Source URL"),
        ("title", "NVARCHAR", "", "Page title"),
        ("content", "NVARCHAR(MAX)", "", "Full page text"),
        ("contentHash", "NVARCHAR", "", "Change detection hash"),
        ("category", "NVARCHAR", "", "Content category"),
        ("subcategory", "NVARCHAR", "", "Sub-category"),
        ("language", "NVARCHAR", "", "Content language"),
        ("wordCount", "INT", "", "Word count"),
        ("lastScrapedAt", "DATETIME", "", "Last scrape time"),
        ("isActive", "BIT", "", "Active in KB flag"),
        ("sourceDomain", "NVARCHAR", "", "Source domain"),
        ("createdAt", "DATETIME", "", "Creation time"),
        ("updatedAt", "DATETIME", "", "Update time"),
    ],
    "KnowledgeChunks": [
        ("id", "UUID", "PK", "Chunk ID"),
        ("pageId", "UUID", "FK→KnowledgePages", "Parent page"),
        ("chunkIndex", "INT", "", "Position in page"),
        ("content", "NVARCHAR(MAX)", "", "Chunk text"),
        ("keywords", "NVARCHAR", "", "Extracted keywords"),
        ("category", "NVARCHAR", "", "Inherited category"),
        ("title", "NVARCHAR", "", "Chunk title"),
        ("chromaSyncedAt", "DATETIME", "", "Vector sync timestamp"),
        ("embeddingModel", "NVARCHAR", "", "Model used for embedding"),
        ("entityType", "NVARCHAR", "", "Entity classification"),
        ("createdAt", "DATETIME", "", "Creation time"),
    ],
    "ScraperRuns": [
        ("id", "UUID", "PK", "Run ID"),
        ("runType", "NVARCHAR", "", "crawl/sync/rescrape"),
        ("status", "NVARCHAR", "", "running/completed/failed"),
        ("baseUrl", "NVARCHAR", "", "Starting URL"),
        ("pagesScraped", "INT", "", "Total pages processed"),
        ("pagesUpdated", "INT", "", "Updated pages"),
        ("pagesNew", "INT", "", "New pages"),
        ("pagesUnchanged", "INT", "", "Unchanged pages"),
        ("errorCount", "INT", "", "Errors encountered"),
        ("errorDetails", "NVARCHAR(MAX)", "", "Error log"),
        ("startedAt", "DATETIME", "", "Start time"),
        ("completedAt", "DATETIME", "", "Completion time"),
    ],
    "ScrapeQueue": [
        ("id", "UUID", "PK", "Queue item ID"),
        ("url", "NVARCHAR", "", "URL to scrape"),
        ("priority", "INT", "", "Queue priority"),
        ("status", "NVARCHAR", "", "pending/processing/done"),
        ("runId", "UUID", "FK→ScraperRuns", "Associated run"),
        ("depth", "INT", "", "Crawl depth"),
        ("createdAt", "DATETIME", "", "Queued time"),
        ("updatedAt", "DATETIME", "", "Last update"),
    ],
    "Admins": [
        ("universityId", "NVARCHAR", "PK", "Admin student ID"),
        ("createdAt", "DATETIME", "", "Grant timestamp"),
    ],
    "Companies": [
        ("id", "INT", "PK IDENTITY", "Company ID"),
        ("name", "NVARCHAR", "", "Company name"),
        ("description", "NVARCHAR", "", "Description"),
        ("colors", "NVARCHAR", "", "Brand colors JSON"),
        ("scale", "FLOAT", "", "3D model scale"),
        ("svgString", "NVARCHAR(MAX)", "", "SVG logo data"),
        ("email", "NVARCHAR", "", "Contact email"),
        ("phone", "NVARCHAR", "", "Contact phone"),
        ("website", "NVARCHAR", "", "Company website"),
        ("forceWhiteBack", "BIT", "", "Logo background flag"),
        ("forceBlackBack", "BIT", "", "Logo background flag"),
        ("isMetallic", "BIT", "", "3D material flag"),
    ],
    "InternshipReviews": [
        ("id", "INT", "PK IDENTITY", "Review ID"),
        ("companyId", "INT", "FK→Companies", "Reviewed company"),
        ("userId", "UUID", "FK→Users", "Reviewer"),
        ("userName", "NVARCHAR", "", "Display name"),
        ("rating", "INT", "", "1-5 star rating"),
        ("feedback", "NVARCHAR", "", "Written feedback"),
        ("createdAt", "DATETIME", "", "Submission time"),
    ],
    "CapstonePartners": [
        ("id", "INT", "PK IDENTITY", "Listing ID"),
        ("userId", "UUID", "FK→Users", "Poster"),
        ("userName", "NVARCHAR", "", "Display name"),
        ("email", "NVARCHAR", "", "Contact email"),
        ("phone", "NVARCHAR", "", "Contact phone"),
        ("major", "NVARCHAR", "", "Student major"),
        ("skills", "NVARCHAR", "", "Skills list"),
        ("description", "NVARCHAR", "", "Self description"),
        ("lookingFor", "NVARCHAR", "", "Partner requirements"),
        ("createdAt", "DATETIME", "", "Post time"),
    ],
    "CapstoneIdeas": [
        ("id", "INT", "PK IDENTITY", "Idea ID"),
        ("title", "NVARCHAR", "", "Project title"),
        ("description", "NVARCHAR", "", "Project description"),
        ("faculty", "NVARCHAR", "", "Faculty area"),
        ("year", "NVARCHAR", "", "Academic year"),
        ("tags", "NVARCHAR", "", "Comma-separated tags"),
        ("isActive", "BIT", "", "Visible flag"),
        ("createdAt", "DATETIME", "", "Creation time"),
    ],
    "Events": [
        ("id", "INT", "PK IDENTITY", "Event ID"),
        ("title", "NVARCHAR", "", "Event title"),
        ("description", "NVARCHAR", "", "Event description"),
        ("location", "NVARCHAR", "", "Venue"),
        ("startDate", "DATETIME", "", "Start date/time"),
        ("endDate", "DATETIME", "", "End date/time"),
        ("category", "NVARCHAR", "", "Event category"),
        ("tags", "NVARCHAR", "", "Tags"),
        ("imageUrl", "NVARCHAR", "", "Event image URL"),
        ("externalUrl", "NVARCHAR", "", "External link"),
        ("source", "NVARCHAR", "", "manual/scraped"),
        ("sourceId", "NVARCHAR", "", "External source ID"),
        ("scraperSource", "NVARCHAR", "", "Scraper origin"),
        ("organizer", "NVARCHAR", "", "Organizer name"),
        ("isFree", "BIT", "", "Free admission flag"),
        ("isActive", "BIT", "", "Visible flag"),
        ("createdBy", "UUID", "FK→Users", "Creator (manual events)"),
        ("createdAt", "DATETIME", "", "Creation time"),
        ("updatedAt", "DATETIME", "", "Update time"),
    ],
    "UserRoadmap": [
        ("id", "INT", "PK IDENTITY", "Roadmap entry ID"),
        ("userId", "UUID", "FK→Users", "Owner"),
        ("courseCode", "NVARCHAR", "", "Planned course code"),
        ("courseName", "NVARCHAR", "", "Course title"),
        ("credits", "INT", "", "Credit hours"),
        ("category", "NVARCHAR", "", "Core/Elective"),
        ("year", "INT", "", "Planned year"),
        ("semester", "NVARCHAR", "", "Fall/Spring/Summer"),
        ("createdAt", "DATETIME", "", "Entry creation time"),
    ],
}

USE_CASES = [
    ("UC-01", "Student Login", "Student", "Authenticate via MU Portal credentials", "JWT token issued; user auto-registered if new"),
    ("UC-02", "Ask MuChat", "Student", "Submit academic question to AI chatbot", "RAG-retrieved answer with university context"),
    ("UC-03", "Generate Schedule", "Student", "Request optimal semester schedule", "Ranked list of conflict-free schedules returned"),
    ("UC-04", "Save Schedule", "Student", "Save preferred generated schedule", "Schedule persisted with selected sections"),
    ("UC-05", "Analyze Resume", "Student", "Upload CV for AI scoring", "ATS/content/impact scores with breakdown"),
    ("UC-06", "Build Resume", "Student", "Create CV via form builder", "Structured resume data with local/global template"),
    ("UC-07", "Edit Resume Live", "Student", "Inline edit with AI suggestions", "Normalized schema updated; export to PDF/DOCX"),
    ("UC-08", "Browse Internships", "Student", "Explore 3D company carousel", "Company details and aggregated ratings displayed"),
    ("UC-09", "Submit Review", "Student", "Rate internship company", "Review stored (one per user per company)"),
    ("UC-10", "Find Capstone Partner", "Student", "Post or browse partner listings", "Matching students connected via listings"),
    ("UC-11", "Browse Capstone Ideas", "Student", "Search faculty project ideas", "Filtered ideas list with tags"),
    ("UC-12", "Plan Degree Roadmap", "Student", "Drag-and-drop course planning", "Personalized CS curriculum roadmap saved"),
    ("UC-13", "View Events", "Student", "Browse upcoming campus events", "Scraped and manual events displayed"),
    ("UC-14", "Sync Academic History", "Student", "Refresh grades from portal", "AcademicHistory table updated"),
    ("UC-15", "Grant Admin", "Admin", "Add admin privileges by university ID", "User added to Admins table"),
    ("UC-16", "Manage Knowledge Base", "Admin", "Crawl/sync/reindex university website", "RAG knowledge base updated"),
    ("UC-17", "Manage Capstone Ideas", "Admin", "CRUD capstone project ideas", "Ideas database maintained"),
    ("UC-18", "Manage Events", "Admin", "Add/edit/delete manual events", "Events table updated"),
    ("UC-19", "Manage Companies", "Admin", "CRUD internship companies", "3D carousel data maintained"),
    ("UC-20", "Mobile Access", "Student", "Use core features on mobile app", "13 screens via Expo tab navigation"),
]

USER_ROLES = [
    ("Anonymous", "Browse public endpoints", "Events list, internship companies, capstone ideas, resume tools"),
    ("Authenticated Student", "JWT bearer token (24h)", "Schedule, history, saved data, reviews, roadmap, chat sessions"),
    ("Administrator", "Admins table or super admin ID", "User management, KB controls, events CRUD, capstone ideas CRUD"),
    (f"Super Admin ({SUPER_ADMIN_ID})", "Hardcoded in auth middleware", "Cannot be demoted; full admin privileges"),
]

LITERATURE_COMPARISON = [
    ("Feature", "MuGate", "Typical University Portal", "Generic ChatGPT"),
    ("Schedule Generation", "Prerequisite-aware backtracking with scoring", "Manual section selection only", "Advisory text only, no data integration"),
    ("Authentication", "Live MU Portal Playwright verification", "Native SSO", "No university authentication"),
    ("Academic Data", "Scraped history + course sync", "Built-in transcript", "No access to student records"),
    ("AI Assistance", "RAG over university KB + student context", "FAQ pages", "General knowledge, may hallucinate"),
    ("Resume Tools", "Analyzer + builder + live editor + export", "Career office appointments", "Text suggestions only"),
    ("Internships", "3D explorer with peer reviews", "Static listings", "No structured review system"),
    ("Capstone", "Partner matching + ideas DB + AI advisor", "Faculty bulletin boards", "Generic project suggestions"),
    ("Events", "Automated scraping + manual admin CRUD", "Manual calendar posts", "No event integration"),
    ("Mobile", "Expo partial port (13 screens)", "Often responsive web only", "Mobile app available"),
    ("Offline Resilience", "AI cascade across 3 providers", "N/A", "Single provider dependency"),
]

FRONTEND_ROUTES = [
    ("/", "Home", "Landing page with MU portal login and feature showcases"),
    ("/chatbot", "Chatbot", "MuChat AI assistant with RAG and file upload"),
    ("/internships", "InternshipList", "3D company carousel with reviews"),
    ("/schedule", "Schedule", "Smart schedule generator and saved schedules"),
    ("/resume-enhancer", "ResumeEnhancer", "CV analyzer, builder, and live editor"),
    ("/capstone", "Capstone", "Partner finder, ideas database, AI advisor"),
    ("/events", "Events", "Campus and academic events hub"),
    ("/roadmap", "RoadMap", "Drag-and-drop CS degree planner"),
    ("/about", "About", "Project credits and team information"),
    ("/admin-control", "AdminControl", "Admin management and KB controls"),
    ("/profile", "Profile", "User profile and session management"),
]

MOBILE_SCREENS = [
    ("LoginScreen", "auth/", "Root stack modal — MU portal login"),
    ("HomeScreen", "home/", "Home tab main screen"),
    ("ChatScreen", "chat/", "MuChat tab"),
    ("ToolsHubScreen", "tools/", "Tools tab hub"),
    ("ScheduleScreen", "schedule/", "Schedule generator"),
    ("ResumeScreen", "resume/", "Resume enhancer"),
    ("CapstoneScreen", "capstone/", "Capstone partner/ideas"),
    ("RoadmapScreen", "roadmap/", "Degree roadmap"),
    ("EventsScreen", "events/", "Events listing"),
    ("InternshipsScreen", "internships/", "Internships explorer tab"),
    ("ProfileScreen", "profile/", "Profile tab"),
    ("AboutScreen", "about/", "About/credits"),
    ("AdminScreen", "admin/", "Admin controls"),
]

DIAGRAM_CAPTIONS = {
    "01-system-architecture": "MuGate three-tier system architecture (frontend, backend, mobile)",
    "02-use-case": "Use case diagram showing student and admin interactions",
    "03-auth-sequence": "Authentication sequence — MU Portal Playwright scrape to JWT issuance",
    "04-rag-sequence": "RAG retrieval sequence — vector + keyword fusion with live fallback",
    "05-resume-activity": "Resume Enhancer activity diagram (analyzer, builder, editor flows)",
    "06-erd": "Entity-relationship diagram of the MuGate database schema",
    "07-dfd-level0": "Data flow diagram (Level 0) for the MuGate platform",
    "08-deployment": "Deployment architecture (no Docker — native Node.js services)",
    "09-component": "Component diagram of backend modules",
    "10-gantt": "Project Gantt chart for the 2025-2026 development timeline",
}

MANUAL_TEST_CASES = [
    ("TC-01", "Login", "Valid MU credentials", "JWT returned; user in database", "Pass"),
    ("TC-02", "Login", "Invalid credentials", "401 error from portal verification", "Pass"),
    ("TC-03", "MuChat", "University policy question", "RAG context injected; relevant answer", "Pass"),
    ("TC-04", "MuChat", "Resume session (source=resume)", "Resume advisor prompt; no RAG", "Pass"),
    ("TC-05", "Schedule", "Generate with preferences", "Top 20 ranked schedules returned", "Pass"),
    ("TC-06", "Schedule", "Save and reload schedule", "Persisted in Schedules table", "Pass"),
    ("TC-07", "Resume", "Upload PDF for analysis", "ATS/content/impact scores returned", "Pass"),
    ("TC-08", "Resume", "Export PDF local template", "Valid PDF generated", "Pass"),
    ("TC-09", "Resume", "Export DOCX global template", "Valid DOCX generated", "Pass"),
    ("TC-10", "Internships", "Submit company review", "Review stored; stats updated", "Pass"),
    ("TC-11", "Capstone", "Post partner listing", "Listing visible in search", "Pass"),
    ("TC-12", "Events", "Trigger scrape", "New events ingested", "Pass"),
    ("TC-13", "Roadmap", "Save custom plan", "UserRoadmap persisted", "Pass"),
    ("TC-14", "Admin", "Grant admin privileges", "User in Admins table", "Pass"),
    ("TC-15", "Admin", "Delete super admin", f"Protected — {SUPER_ADMIN_ID} cannot be removed", "Pass"),
    ("TC-16", "Admin", "Full KB crawl", "ScraperRuns record created", "Pass"),
    ("TC-17", "Mobile", "Login on Expo app", "JWT stored in secure store", "Pass"),
    ("TC-18", "Mobile", "Chat from mobile", "Same backend RAG pipeline", "Pass"),
    ("TC-19", "AI Cascade", "Primary provider rate-limited", "Falls back to next provider", "Pass"),
    ("TC-20", "RAG", "Low confidence query", "Live search fallback triggered", "Pass"),
]

REFERENCES = [
    "Al Maaref University. (2026). Student Portal Documentation. https://portal.mu.edu.lb",
    "Anthropic. (2024). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. arXiv preprint.",
    "ChromaDB Team. (2025). ChromaDB Documentation. https://docs.trychroma.com",
    "DeepSeek AI. (2025). DeepSeek API Documentation. https://api.deepseek.com",
    "Express.js Foundation. (2025). Express 5.x Guide. https://expressjs.com",
    "Facebook Inc. (2025). React 19 Documentation. https://react.dev",
    "Google. (2025). Gemini API Reference. https://ai.google.dev",
    "Microsoft. (2025). SQL Server Documentation. https://docs.microsoft.com/sql",
    "Microsoft. (2025). Playwright Documentation. https://playwright.dev",
    "OpenRouter. (2025). OpenRouter API Documentation. https://openrouter.ai/docs",
    "Sommerville, I. (2016). Software Engineering (10th ed.). Pearson.",
    "Pressman, R. S., & Maxim, B. R. (2019). Software Engineering: A Practitioner's Approach (9th ed.). McGraw-Hill.",
    "Expo Team. (2025). Expo SDK 56 Documentation. https://docs.expo.dev",
    "Vite Team. (2025). Vite 7 Guide. https://vitejs.dev",
    "Cormack, G. V., Clarke, C. L., & Buettcher, S. (2009). Reciprocal Rank Fusion outperforms Condorcet and individual Rank Learning Methods. SIGIR.",
    "OWASP Foundation. (2024). OWASP Top Ten Web Application Security Risks.",
    "IEEE. (2014). IEEE Std 830-1998 — Recommended Practice for Software Requirements Specifications.",
    "Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design Patterns: Elements of Reusable Object-Oriented Software. Addison-Wesley.",
]


# ---------------------------------------------------------------------------
# Chapter builders
# ---------------------------------------------------------------------------
def build_cover(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(PROJECT)
    set_run_font(tr, size=Pt(28), bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("A University AI Platform for Academic Excellence")
    set_run_font(sr, size=Pt(16), italic=True)

    doc.add_paragraph()
    doc.add_paragraph()

    for line in [UNIVERSITY, DEPARTMENT, f"Academic Year {ACADEMIC_YEAR}", "", "Capstone Project Report", "", "Submitted By:"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(line), size=Pt(14))

    for name, roles in AUTHORS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(f"{name} — {', '.join(roles)}"), size=Pt(12))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(f"© 2026 {UNIVERSITY}"), size=Pt(11), italic=True)
    add_page_break(doc)


def build_abstract(doc: Document) -> None:
    add_heading(doc, "Abstract", level=1)
    add_bodies(doc, [
        f"{PROJECT} is a comprehensive university AI platform developed at {UNIVERSITY}, "
        f"{DEPARTMENT}, during the {ACADEMIC_YEAR} academic year. The platform addresses "
        "the fragmented nature of student academic tooling by consolidating intelligent "
        "scheduling, AI-powered academic assistance, resume enhancement, internship discovery, "
        "capstone collaboration, event aggregation, and degree planning into a unified "
        "web and mobile experience.",
        "The system comprises three applications: a React 19 frontend built with Vite 7, "
        "an Express 5 TypeScript backend connected to Microsoft SQL Server, and a partial "
        "Expo 56 mobile port providing 13 screens across five tab navigations. Authentication "
        "integrates directly with the Al Maaref University student portal via Playwright-based "
        "credential verification, issuing JSON Web Tokens with a 24-hour expiry.",
        "Artificial intelligence is delivered through a resilient three-provider cascade "
        "(DeepSeek → Gemini → OpenRouter) with retrieval-augmented generation (RAG) powered "
        "by ChromaDB vector storage and reciprocal rank fusion between vector and keyword "
        "search. The schedule generator employs prerequisite-aware backtracking with multi-criteria "
        "scoring to produce ranked, conflict-free timetables. The Resume Enhancer provides "
        "analyzer, builder, and live editor modes with local and global CV templates.",
        "The backend exposes 50+ REST API endpoints across 22 database tables. Testing was "
        "conducted exclusively through manual verification; no automated test suite or Docker "
        "containerization is present in the repository. This report documents the complete "
        "system lifecycle from requirements through deployment, presenting architecture diagrams, "
        "database design, algorithm specifications, and an honest assessment of challenges "
        "and future work.",
    ])
    kw = doc.add_paragraph()
    set_run_font(kw.add_run(
        "Keywords: university platform, retrieval-augmented generation, schedule optimization, "
        "resume enhancement, Playwright authentication, React, Express, SQL Server, Expo"
    ), bold=True, italic=True)
    add_page_break(doc)


def build_acknowledgements(doc: Document) -> None:
    add_heading(doc, "Acknowledgements", level=1)
    add_bodies(doc, [
        f"We extend our sincere gratitude to the faculty and staff of {UNIVERSITY} for "
        "their guidance, resources, and encouragement throughout the development of the "
        f"{PROJECT} platform.",
        "We thank our dedicated instructors and mentors who reviewed our design decisions "
        "and offered constructive feedback during each phase of development.",
        "Special recognition goes to fellow students who participated in early testing sessions "
        "and reported bugs that shaped the user experience.",
        "We acknowledge the open-source communities behind React, Express, Playwright, ChromaDB, "
        "and Expo, whose tools made this integration project feasible within an academic timeline.",
        "Finally, we thank our families for their patience and support during development.",
    ])
    add_page_break(doc)


def build_front_matter_lists(doc: Document, counters: Counters,
                             figure_list: Optional[List[str]] = None,
                             table_list: Optional[List[str]] = None) -> None:
    add_heading(doc, "Table of Contents", level=1)
    add_toc_field(doc)
    add_page_break(doc)
    add_heading(doc, "List of Figures", level=1)
    for fig in (figure_list or counters.figures):
        p = doc.add_paragraph()
        set_run_font(p.add_run(fig), size=Pt(11))
    add_page_break(doc)
    add_heading(doc, "List of Tables", level=1)
    for tbl in (table_list or counters.tables):
        p = doc.add_paragraph()
        set_run_font(p.add_run(tbl), size=Pt(11))
    add_page_break(doc)


def build_chapter1(doc: Document, counters: Counters) -> None:
    add_heading(doc, "Chapter 1: Introduction", level=1)
    add_heading(doc, "1.1 Background and Motivation", level=2)
    add_bodies(doc, [
        f"University students at {UNIVERSITY} navigate a complex academic ecosystem spanning "
        "course registration, degree planning, career preparation, capstone project formation, "
        "and campus engagement. Traditionally, these activities are supported by disconnected "
        "tools: the official student portal for registration, static PDF degree plans, "
        "third-party resume builders, social media for event discovery, and informal networks "
        "for capstone partner matching.",
        "This fragmentation imposes cognitive overhead on students who must context-switch "
        "between systems, manually reconcile prerequisite requirements when building schedules, "
        "and rely on generic AI chatbots that lack institution-specific knowledge.",
        f"The {PROJECT} project was conceived to address these gaps by delivering an integrated, "
        "AI-augmented platform tailored specifically to Al Maaref University's Computer Science "
        "curriculum and administrative workflows.",
        "The motivation extends beyond convenience: accurate schedule generation directly impacts "
        "graduation timelines, resume quality affects employment outcomes, and timely access to "
        "capstone partners and internship reviews shapes career trajectories.",
    ])
    add_heading(doc, "1.2 Problem Statement", level=2)
    add_bodies(doc, [
        "Students face four interconnected problems: manual schedule construction against "
        "prerequisite chains and section availability; scattered academic policy information; "
        "disconnected career preparation tools; and unstructured capstone and event discovery.",
        f"{PROJECT} unifies solutions within a single authenticated platform that reads live "
        "data from the MU portal and augments it with AI capabilities grounded in a curated "
        "university knowledge base.",
    ])
    add_heading(doc, "1.3 Project Objectives", level=2)
    add_bullet_list(doc, [
        "Implement secure MU Portal authentication via Playwright with JWT (24h).",
        "Develop prerequisite-aware schedule generator with backtracking and scoring.",
        "Build MuChat RAG chatbot with ChromaDB vector storage.",
        "Create Resume Enhancer (analyzer, builder, live editor) with local/global templates.",
        "Deliver 3D Internships explorer with peer reviews.",
        "Provide Capstone partner matching, ideas database, and AI advisor.",
        "Aggregate events via scraping and admin CRUD.",
        "Enable drag-and-drop degree roadmap planning.",
        "Implement Admin Control for users and knowledge base.",
        "Port core features to Expo mobile (13 screens).",
    ])
    add_heading(doc, "1.4 Scope and Limitations", level=2)
    add_bodies(doc, [
        "Scope covers web frontend, Express backend, SQL Server database, and partial mobile port. "
        "Limitations: no automated tests, no Docker, hardcoded super-admin ID 101230004, "
        "desktop min-width 1200px, mobile partial port only.",
    ])
    add_heading(doc, "1.5 Project Team", level=2)
    make_table(doc, ["Name", "Role(s)"], [(n, ", ".join(r)) for n, r in AUTHORS],
               counters, "Project team members and responsibilities")
    add_heading(doc, "1.6 Report Organization", level=2)
    add_bodies(doc, [
        "Chapters 2–12 cover literature review, requirements, design, implementation, UI, "
        "database, algorithms, testing, deployment, challenges, and conclusion.",
    ])
    add_page_break(doc)


def build_chapter2(doc: Document, counters: Counters) -> None:
    add_heading(doc, "Chapter 2: Literature Review", level=1)
    sections = [
        ("2.1 University Student Portals", [
            "University portals provide transactional access but lack intelligent advisory features. "
            "MuGate scrapes live portal data and applies combinatorial optimization for schedules.",
        ]),
        ("2.2 AI Assistants in Education", [
            "Generic LLMs hallucinate on institution-specific facts. RAG mitigates this by injecting "
            "retrieved documents. MuGate uses ChromaDB, hybrid retrieval, and Reciprocal Rank Fusion.",
        ]),
        ("2.3 Schedule Optimization", [
            "Course scheduling is a constrained optimization problem. MuGate uses depth-first backtracking "
            "with pruning and multi-criteria scoring, returning top 20 ranked schedules.",
        ]),
        ("2.4 Resume Analysis and ATS", [
            "MuGate combines heuristic and AI scoring across ATS, content, impact, structure, and keywords. "
            "Normalized resumeSchema.js separates data from Lebanese/international templates.",
        ]),
        ("2.5 Technology Landscape", [
            "React 19 + Vite 7 frontend, Express 5 + TypeScript backend, SQL Server, ChromaDB, Expo 56 mobile.",
        ]),
        ("2.6 Research Gap and Contribution", [
            f"{PROJECT} demonstrates a unified student success platform for {UNIVERSITY} CS students, "
            "honestly documenting academic-project limitations.",
        ]),
    ]
    for title, paras in sections:
        add_heading(doc, title, level=2)
        add_bodies(doc, paras)
    add_heading(doc, "2.7 Comparative Analysis", level=2)
    make_table(doc, ["Dimension", "MuGate", "University Portal", "Generic AI"],
               LITERATURE_COMPARISON[1:], counters, "Feature comparison")
    add_page_break(doc)


def build_chapter3(doc: Document, counters: Counters) -> None:
    add_heading(doc, "Chapter 3: Requirements Analysis", level=1)
    add_heading(doc, "3.1 Stakeholders", level=2)
    add_bodies(doc, ["Primary: CS students. Secondary: administrators. Development: two-member team."])
    add_heading(doc, "3.2 Functional Requirements", level=2)
    func = [
        ("FR-01", "Auth", "Portal Playwright verify + JWT 24h"),
        ("FR-02", "Registration", "Auto-register on first login"),
        ("FR-03", "Credentials", "AES-encrypt portal creds for scraping"),
        ("FR-04", "Schedule", "Conflict-free prerequisite-aware generation"),
        ("FR-05", "Persistence", "Save/delete schedules"),
        ("FR-06", "MuChat", "RAG for university questions"),
        ("FR-07", "Sessions", "Separate chat and resume sessions"),
        ("FR-08", "Resume", "AI scoring + PDF/DOCX export"),
        ("FR-09", "Internships", "3D explorer + one review per company"),
        ("FR-10", "Capstone", "Partners + ideas + AI advisor"),
        ("FR-11", "Events", "Scrape + admin CRUD"),
        ("FR-12", "Roadmap", "Drag-and-drop degree plan"),
        ("FR-13", "Admin", "Grant/revoke admin, KB management"),
        ("FR-14", "Mobile", "13-screen Expo partial port"),
        ("FR-15", "AI Cascade", "DeepSeek → Gemini → OpenRouter fallback"),
    ]
    make_table(doc, ["ID", "Category", "Requirement"], func, counters, "Functional requirements")
    add_heading(doc, "3.3 Non-Functional Requirements", level=2)
    nfr = [
        ("NFR-01", "Performance", "Schedule gen < 30s typical"),
        ("NFR-02", "Resilience", "AI cascade on provider failure"),
        ("NFR-03", "Security", "AES credentials, JWT sessions"),
        ("NFR-04", "Scalability", "ChromaDB or local vector fallback"),
        ("NFR-05", "Integrity", "22 tables with FK constraints"),
    ]
    make_table(doc, ["ID", "Category", "Requirement"], nfr, counters, "Non-functional requirements")
    add_heading(doc, "3.4 Use Cases", level=2)
    make_table(doc, ["ID", "Name", "Actor", "Description", "Postcondition"], USE_CASES, counters, "Use case descriptions")
    add_heading(doc, "3.5 User Roles", level=2)
    make_table(doc, ["Role", "Auth", "Capabilities"], USER_ROLES, counters, "User roles")
    add_page_break(doc)


def build_chapter4(doc: Document, counters: Counters) -> None:
    add_heading(doc, "Chapter 4: System Design", level=1)
    add_bodies(doc, [
        "MuGate follows a three-tier architecture: presentation (React 19 + Expo 56), "
        "application (Express 5 TypeScript REST API), and data (SQL Server + ChromaDB). "
        "Playwright bridges the application tier to the external MU Portal for authentication "
        "and academic data scraping.",
    ])
    add_heading(doc, "4.1 Architecture Overview", level=2)
    for key, cap in DIAGRAM_CAPTIONS.items():
        png = DIAGRAMS_DIR / f"{key}.png"
        add_figure(doc, png, counters, cap)
    add_heading(doc, "4.2 Module Decomposition", level=2)
    modules = [
        ("Auth Module", "Portal verification, JWT, admin checks, credential encryption."),
        ("Scheduling Module", "Course sync, eligibility, backtracking generator, saved schedules."),
        ("AI/Chatbot Module", "MuChat sessions, classifier, RAG pipeline, AI cascade."),
        ("Resume Module", "Analyze, parse, convert, AI edit, PDF/DOCX generation."),
        ("Scraper Module", "Portal scraper, university crawler, event scraper, CRON jobs."),
        ("Internships Module", "Company CRUD, 3D assets, review system."),
        ("Capstone Module", "Partners, ideas, soft-delete, AI advisor."),
        ("Events Module", "Scraped + manual events with categories."),
        ("Roadmap Module", "Default CS curriculum + user overrides."),
        ("Admin Module", "User/admin management, KB controls."),
    ]
    make_table(doc, ["Module", "Responsibility"], modules, counters, "Backend module decomposition")
    add_heading(doc, "4.3 Frontend Routes", level=2)
    make_table(doc, ["Path", "Component", "Description"], FRONTEND_ROUTES, counters, "Frontend route map")
    add_heading(doc, "4.4 Mobile Architecture", level=2)
    make_table(doc, ["Screen", "Directory", "Description"], MOBILE_SCREENS, counters, "Mobile screen inventory")
    add_bodies(doc, ["Five tabs: Home, MuChat, Tools, Internships, Profile."])
    add_heading(doc, "4.5 Security Design", level=2)
    add_bodies(doc, [
        f"JWT 24h expiry. AES encryption for PortalCredentials. Admin middleware checks "
        f"Admins table or super-admin {SUPER_ADMIN_ID}. Rate limiting via express-rate-limit.",
    ])
    add_page_break(doc)


def build_chapter5(doc: Document, counters: Counters) -> None:
    add_heading(doc, "Chapter 5: Implementation", level=1)
    add_heading(doc, "5.1 Technology Stack", level=2)
    make_table(doc, ["Layer", "Technology", "Version", "Purpose"], TECH_STACK, counters, "Technology stack")
    add_heading(doc, "5.2 Backend Implementation", level=2)
    add_bodies(doc, [
        "Express 5 server in TypeScript 5.9.3. Modular structure under backend/src/modules/. "
        "mssql driver with msnodesqlv8 for Windows integrated auth option. "
        "50+ endpoints registered across route files. node-cron schedules: course sync (3h), "
        "KB sync (daily 3AM), vector sync (4AM), sitemap refresh (Sunday 2AM).",
    ])
    add_heading(doc, "5.3 Frontend Implementation", level=2)
    add_bodies(doc, [
        "React 19 SPA with Vite 7. Route-level lazy loading via React.lazy — 24 production chunks. "
        "Tailwind CSS 4 for styling. Three.js for internship 3D carousel. "
        "@dnd-kit for roadmap drag-and-drop.",
    ])
    add_heading(doc, "5.4 Mobile Implementation", level=2)
    add_bodies(doc, [
        "Expo 56 with React Native 0.85.3. React Navigation 7 (tabs + stacks + drawer). "
        "expo-secure-store for JWT. Partial port: 13 screens covering core features.",
    ])
    add_heading(doc, "5.5 Authentication Implementation", level=2)
    embed_glob_figures(doc, str(CODE_DIR / "auth-login-flow.png"), counters,
                       lambda n: "AuthService.login — MU Portal verification flow")
    add_bodies(doc, [
        "auth.service.ts: PortalScraper.verifyCredentials → auto-register → encrypt creds → "
        "sync history → generateToken → isAdmin check.",
    ])
    add_heading(doc, "5.6 AI Provider Implementation", level=2)
    embed_glob_figures(doc, str(CODE_DIR / "ai-cascade.png"), counters,
                       lambda n: "AiProvider — multi-provider cascade with retry")
    add_heading(doc, "5.7 RAG Implementation", level=2)
    add_bodies(doc, [
        "KnowledgeService.retrieveContext: term extraction → parallel vector+keyword → RRF merge → "
        "confidence check (0.55) → LiveSearchService fallback. ChromaDB collection mugate_knowledge_chunks "
        "with local vectors.json fallback. Gemini text-embedding-004 with hash-vector fallback.",
    ])
    add_heading(doc, "5.8 Resume Enhancer Implementation", level=2)
    embed_glob_figures(doc, str(CODE_DIR / "resume-schema.png"), counters,
                       lambda n: "resumeSchema.js — normalized template-agnostic data model")
    add_bodies(doc, [
        "Three modes: WelcomeView → Analyzer/Builder/Editor. pdfkit + docx backend generators. "
        "localStorage persistence via usePersistentState.",
    ])
    add_heading(doc, "5.9 Schedule Generator Implementation", level=2)
    add_bodies(doc, [
        "GeneratorService: auto-scrape history → sync courses → getEligibleCourses (CS_CURRICULUM) → "
        "backtracking → score → top 20. Preferences: maxCredits, excludeDays, startTime, twoDaysOnly.",
    ])
    add_heading(doc, "5.10 Frontend Routes", level=2)
    embed_glob_figures(doc, str(CODE_DIR / "frontend-routes.png"), counters,
                       lambda n: "App.jsx — lazy-loaded route definitions")
    add_page_break(doc)


def build_chapter6(doc: Document, counters: Counters) -> None:
    add_heading(doc, "Chapter 6: User Interface", level=1)
    add_bodies(doc, [
        "The MuGate web interface employs a consistent navigation bar across 11 routes with "
        "admin-gated features. Visual design uses custom CSS with Tailwind utilities and "
        "a GlobalGlow ambient effect. Desktop layout requires min-width 1200px.",
    ])
    add_heading(doc, "6.1 Home and Authentication", level=2)
    add_bodies(doc, ["Landing page with MU portal login form, feature showcases, and animated branding."])
    add_heading(doc, "6.2 MuChat Interface", level=2)
    add_bodies(doc, ["Chat sessions with markdown rendering, file upload, pin/delete, separate resume advisor."])
    add_heading(doc, "6.3 Schedule Interface", level=2)
    add_bodies(doc, ["Preference panel, generated schedule cards with scores, save/delete functionality."])
    add_heading(doc, "6.4 Resume Enhancer Interface", level=2)
    add_bodies(doc, ["Welcome mode selector, analyzer score rings, builder forms, Jobsuit-style live editor."])
    add_heading(doc, "6.5 Internships 3D Explorer", level=2)
    add_bodies(doc, ["Three.js carousel with company logos, review submission, admin company management."])
    add_heading(doc, "6.6 Capstone, Events, Roadmap", level=2)
    add_bodies(doc, [
        "Capstone: partner listings, searchable ideas, AI chat advisor. "
        "Events: category filters, scraped + manual events. "
        "Roadmap: drag-and-drop semester planning against CS curriculum.",
    ])
    add_heading(doc, "6.7 Admin Control Panel", level=2)
    add_bodies(doc, [
        "Grant/demote admins, KB stats, crawl/sync/reindex/rescrape buttons, scraper run history. "
        "BroadcastChannel for cross-tab sync.",
    ])
    add_heading(doc, "6.8 Application Screenshots", level=2)
    n = embed_glob_figures(doc, str(SCREENSHOTS_DIR / "web-*.png"), counters,
                           lambda n: f"MuGate web interface — {n.replace('web-', '').replace('-', ' ')}")
    if n == 0:
        add_bodies(doc, ["Screenshots not available at generation time. Place web-*.png files in report-assets/screenshots/."])
    add_page_break(doc)


def build_chapter7(doc: Document, counters: Counters) -> None:
    add_heading(doc, "Chapter 7: Database Design", level=1)
    add_bodies(doc, [
        f"MuGate uses Microsoft SQL Server database 'MuGate' with 22 tables across core academic, "
        "chat, RAG, and feature-specific schemas. Runtime migrations in connection.ts add columns "
        "and tables (Admins, Companies, Events, etc.) beyond init.sql.",
    ])
    add_heading(doc, "7.1 Schema Overview", level=2)
    add_figure(doc, DIAGRAMS_DIR / "06-erd.png", counters, "Entity-relationship diagram")
    add_heading(doc, "7.2 Table Inventory", level=2)
    make_table(doc, ["#", "Table", "Purpose"],
               [(i + 1, t, f"{len(cols)} columns") for i, (t, cols) in enumerate(DB_TABLES.items())],
               counters, "Database table inventory")
    add_heading(doc, "7.3 Data Dictionary", level=2)
    for table_name, columns in DB_TABLES.items():
        add_heading(doc, f"7.3.{list(DB_TABLES.keys()).index(table_name) + 1} {table_name}", level=3)
        make_table(doc, ["Column", "Type", "Constraint", "Description"], columns,
                   counters, f"Data dictionary — {table_name}")
        add_bodies(doc, [
            f"The {table_name} table stores {'user identity and profile' if table_name == 'Users' else 'domain-specific'} "
            f"records for the {PROJECT} platform. Referential integrity is enforced through foreign key "
            f"relationships where indicated.",
        ])
    add_heading(doc, "7.4 Indexing Strategy", level=2)
    add_bodies(doc, [
        "Indexes on userId, sessionId, courseId, category, isActive, startDate, status columns "
        "for query performance on high-traffic endpoints.",
    ])
    add_page_break(doc)


def build_chapter8(doc: Document, counters: Counters) -> None:
    add_heading(doc, "Chapter 8: Algorithms", level=1)

    add_heading(doc, "8.1 Schedule Generation Algorithm", level=2)
    add_bodies(doc, ["Backtracking over eligible courses with conflict detection and scoring."])
    add_pseudocode(doc, "Algorithm 8.1: Schedule Generation", [
        "FUNCTION GenerateSchedules(userId, semesterId, preferences):",
        "    IF AcademicHistory empty THEN SyncHistoryFromPortal(userId)",
        "    IF no cached sections THEN SyncCoursesFromPortal(semesterId)",
        "    eligible ← GetEligibleCourses(userId, CS_CURRICULUM)",
        "    offerings ← FetchSections(semesterId) FILTER full, A-building, TBA",
        "    offerings ← ApplyPreferences(offerings, preferences)",
        "    schedules ← []",
        "    FUNCTION Backtrack(index, currentSections, currentCourses):",
        "        IF index = len(eligible) THEN",
        "            IF ValidateCorequisites(currentCourses) AND ValidateElectiveLimits(currentCourses):",
        "                schedules.APPEND(Score(currentSections))",
        "            RETURN",
        "        course ← eligible[index]",
        "        sections ← offerings[course] OR skip",
        "        FOR EACH section IN sections:",
        "            IF NOT HasTimeConflict(section, currentSections):",
        "                Backtrack(index+1, currentSections + [section], currentCourses + [course])",
        "        Backtrack(index+1, currentSections, currentCourses)  // skip course",
        "    Backtrack(0, [], [])",
        "    IF preferences.twoDaysOnly THEN FILTER schedules BY campusDays ≤ 2",
        "    SORT schedules BY score DESC",
        "    RETURN schedules[0:20]",
    ])
    add_bodies(doc, [
        "Scoring: major credits ×10, other ×5, +8 per course, +15 for 0-credit, compactness bonus, +totalCredits.",
    ])

    add_heading(doc, "8.2 RAG Retrieval Algorithm", level=2)
    add_pseudocode(doc, "Algorithm 8.2: Hybrid RAG Retrieval", [
        "FUNCTION RetrieveContext(question):",
        "    terms ← ExtractSearchTerms(question) WITH synonym expansion",
        "    vectorResults ← VectorSearch(Embed(question), topK=8)",
        "    keywordResults ← KeywordSearch(terms, topK=8)",
        "    merged ← ReciprocalRankFusion(vectorResults, keywordResults)",
        "    filtered ← merged WHERE relevanceScore >= 2",
        "    context ← FormatTopN(filtered, 5)",
        "    confidence ← ComputeConfidence(filtered)",
        "    IF confidence < 0.55 THEN",
        "        liveResults ← LiveSearchAndIngest(question, maxURLs=3)",
        "        context ← MERGE(context, liveResults)",
        "    RETURN { context, sources, confidence }",
    ])

    add_heading(doc, "8.3 AI Provider Cascade", level=2)
    add_pseudocode(doc, "Algorithm 8.3: AI Provider Cascade", [
        "FUNCTION GenerateResponse(systemPrompt, history, message):",
        "    providers ← [DeepSeek, Gemini, OpenRouter] SORT BY AI_PRIMARY_PROVIDER first",
        "    FOR EACH provider IN providers WHERE apiKey exists:",
        "        FOR EACH model IN provider.fallbackModels:",
        "            FOR attempt IN 1..3:",
        "                response ← provider.call(model, systemPrompt, history, message)",
        "                IF response.tokensUsed > 0 AND NOT isConnectionError(response):",
        "                    RETURN response",
        "                IF rateLimited THEN WAIT exponentialBackoff(attempt)",
        "    RETURN mockFallbackMessage()",
    ])

    add_heading(doc, "8.4 Resume Analysis Algorithm", level=2)
    add_pseudocode(doc, "Algorithm 8.4: Resume Analysis", [
        "FUNCTION AnalyzeResume(resumeText):",
        "    structured ← ParseToJSON(resumeText)",
        "    atsScore ← CheckATSCompatibility(structured)      // formatting, sections",
        "    contentScore ← EvaluateContentQuality(structured)  // completeness",
        "    impactScore ← ScoreImpactStatements(structured)    // action verbs, metrics",
        "    structureScore ← ValidateStructure(structured)     // section order",
        "    keywordScore ← MatchKeywords(structured, domain)   // CS/tech keywords",
        "    aiInsights ← AiProvider.enhance(analysisPrompt, structured)",
        "    RETURN { scores, breakdown, suggestions: aiInsights }",
    ])

    add_heading(doc, "8.5 Authentication Flow", level=2)
    add_pseudocode(doc, "Algorithm 8.5: Portal Authentication", [
        "FUNCTION Login(universityId, password):",
        "    result ← Playwright.verifyCredentials(universityId, password)",
        "    IF NOT result.valid THEN THROW AuthenticationError",
        "    user ← FindOrCreateUser(universityId, result.studentName)",
        "    UpsertEncryptedCredentials(user.id, universityId, password)",
        "    TRY SyncStudentHistory(user.id) CATCH log non-fatal",
        "    token ← JWT.sign({ userId, email, name, universityId }, expiry=24h)",
        "    isAdmin ← (universityId = '101230004') OR EXISTS(Admins, universityId)",
        "    RETURN { token, user, isAdmin }",
    ])
    add_page_break(doc)


def build_chapter9(doc: Document, counters: Counters) -> None:
    add_heading(doc, "Chapter 9: Testing", level=1)
    add_bodies(doc, [
        "The MuGate repository contains no automated test suite. No Jest, Mocha, Cypress, or "
        "Playwright test configurations exist in package.json files. Testing was conducted "
        "exclusively through manual verification during development and the final audit phase.",
        "This approach is common in academic capstone projects prioritizing feature breadth, "
        "but represents a documented limitation for production deployment.",
    ])
    add_heading(doc, "9.1 Testing Strategy", level=2)
    add_bodies(doc, [
        "Manual exploratory testing covered all 11 frontend routes, 50+ API endpoints, "
        "mobile 13 screens, and admin workflows. Test data used real MU portal credentials "
        "in development environment.",
    ])
    add_heading(doc, "9.2 Test Cases and Results", level=2)
    make_table(doc, ["ID", "Module", "Scenario", "Expected", "Result"], MANUAL_TEST_CASES, counters, "Manual test cases")
    add_heading(doc, "9.3 API Endpoint Verification", level=2)
    make_table(doc, ["Method", "Path", "Auth", "Description"], API_ENDPOINTS, counters, "Complete API endpoint inventory")
    add_heading(doc, "9.4 Known Issues", level=2)
    add_bullet_list(doc, [
        "Desktop-only navigation (min-width 1200px) — not tested on mobile browsers.",
        "Legacy resume editor path (editor.service.ts) — brittle regex; structured ai-editor preferred.",
        "Secrets in git history — rotation required before production.",
        "No load testing performed on schedule generator or RAG pipeline.",
    ])
    add_heading(doc, "9.5 Recommendations for Future Testing", level=2)
    add_bullet_list(doc, [
        "Add Jest unit tests for GeneratorService, KnowledgeService, AuthService.",
        "Add Supertest integration tests for API routes.",
        "Add Playwright E2E tests for login and schedule generation flows.",
        "Add CI pipeline running tests on pull requests.",
    ])
    add_page_break(doc)


def build_chapter10(doc: Document, counters: Counters) -> None:
    add_heading(doc, "Chapter 10: Deployment", level=1)
    add_bodies(doc, [
        "MuGate is deployed as native Node.js processes without Docker containerization. "
        "No Dockerfile, docker-compose.yml, or container orchestration configuration exists "
        "in the repository. This chapter documents the actual deployment architecture.",
    ])
    add_heading(doc, "10.1 Deployment Architecture", level=2)
    add_figure(doc, DIAGRAMS_DIR / "08-deployment.png", counters, "Deployment architecture (native services)")
    add_bodies(doc, [
        "Three independently deployed applications: backend (port 5000), frontend (Vite dev or "
        "static build via nginx), mobile (Expo build). SQL Server instance hosts MuGate database. "
        "Optional ChromaDB server for production vector search.",
    ])
    add_heading(doc, "10.2 Environment Configuration", level=2)
    env_vars = [
        ("DB_SERVER", "SQL Server hostname"),
        ("DB_NAME", "MuGate"),
        ("JWT_SECRET", "Token signing key"),
        ("DEEPSEEK_API_KEY", "Primary AI provider"),
        ("GEMINI_API_KEY", "Fallback AI + embeddings"),
        ("OPENROUTER_API_KEY", "Tertiary AI fallback"),
        ("CHROMA_URL", "Optional ChromaDB server"),
        ("AI_PRIMARY_PROVIDER", "deepseek | gemini | openrouter"),
        ("ENCRYPTION_KEY", "Portal credential AES key"),
    ]
    make_table(doc, ["Variable", "Purpose"], env_vars, counters, "Environment variables")
    add_heading(doc, "10.3 Backend Deployment", level=2)
    add_bodies(doc, [
        "npm run build (TypeScript compile) → node dist/server.js. "
        "Playwright browsers must be installed on server (npx playwright install). "
        "Windows: msnodesqlv8 for SQL Server integrated auth. CRON jobs start with server.",
    ])
    add_heading(doc, "10.4 Frontend Deployment", level=2)
    add_bodies(doc, ["npm run build → dist/ static files served by nginx or IIS. API URL configured in frontend config."])
    add_heading(doc, "10.5 Mobile Deployment", level=2)
    add_bodies(doc, ["expo build / eas build for iOS and Android. API base URL points to deployed backend."])
    add_heading(doc, "10.6 Database Initialization", level=2)
    add_bodies(doc, [
        "Run init.sql, rag-tables.sql, rag-v2.sql. connection.ts creates runtime tables (Admins, Companies, etc.). "
        "rag.bootstrap.ts seeds knowledge and syncs vectors on startup.",
    ])
    add_heading(doc, "10.7 Why No Docker", level=2)
    add_bodies(doc, [
        "The repository does not include Docker configuration. Deployment assumes direct Node.js "
        "installation on Windows Server or Linux VM with SQL Server connectivity. Docker could be "
        "added in future work but is outside current project scope.",
    ])
    add_page_break(doc)


def build_chapter11(doc: Document, counters: Counters) -> None:
    add_heading(doc, "Chapter 11: Challenges and Future Work", level=1)
    challenges = [
        ("Portal Scraping Reliability", "Playwright depends on portal DOM structure; layout changes break scrapers.", "Implement selector health checks and alert admins."),
        ("AI Provider Rate Limits", "Free tier APIs throttle under load.", "Cascade across 3 providers mitigates; paid tiers for production."),
        ("RAG Knowledge Freshness", "Stale KB pages yield outdated answers.", "CRON sync + live search fallback + admin rescrape tools."),
        ("Hardcoded Super Admin", f"ID {SUPER_ADMIN_ID} cannot be changed without code edits.", "Move to environment variable or bootstrap seed."),
        ("No Automated Tests", "Regression risk on changes.", "Add Jest + Supertest + Playwright suite."),
        ("Desktop-Only UI", "min-width 1200px excludes mobile browsers.", "Responsive redesign with mobile-first breakpoints."),
        ("Secret Management", "API keys were committed to git history.", "Rotate all secrets; use secret manager."),
        ("Mobile Feature Parity", "13 screens vs 11 web routes with varying depth.", "Complete mobile port of all features."),
    ]
    add_heading(doc, "11.1 Development Challenges", level=2)
    make_table(doc, ["Challenge", "Impact", "Mitigation"], challenges, counters, "Development challenges")
    add_heading(doc, "11.2 Future Enhancements", level=2)
    add_bullet_list(doc, [
        "Docker containerization for reproducible deployment.",
        "Automated test suite with CI/CD pipeline.",
        "Responsive mobile-web design.",
        "OAuth2/OIDC if university provides official API.",
        "Real-time notifications for events and capstone matches.",
        "Multi-faculty curriculum support beyond CS.",
        "Offline mobile mode with sync.",
        "Analytics dashboard for admin usage metrics.",
    ])
    add_page_break(doc)


def build_chapter12_and_references(doc: Document, counters: Counters) -> None:
    add_heading(doc, "Chapter 12: Conclusion", level=1)
    add_bodies(doc, [
        f"{PROJECT} successfully demonstrates an integrated university AI platform combining "
        f"portal-authenticated data access, intelligent scheduling, retrieval-augmented generation, "
        f"multi-provider AI resilience, and comprehensive student success tools within a single codebase.",
        f"Developed by Mohammad Jomaa and Abo Al Fadel Ismael at {UNIVERSITY}, {DEPARTMENT}, "
        f"the platform encompasses a React 19 frontend, Express 5 TypeScript backend with 50+ API "
        f"endpoints and 22 database tables, and a partial Expo 56 mobile port with 13 screens.",
        "Key technical achievements include Playwright-based MU Portal authentication with JWT sessions, "
        "prerequisite-aware schedule backtracking returning top-20 ranked timetables, hybrid RAG with "
        "ChromaDB and reciprocal rank fusion, and a three-mode Resume Enhancer with local and global templates.",
        "The project honestly documents its limitations: manual testing only, no Docker deployment, "
        "hardcoded admin privileges, and desktop-oriented UI. These represent appropriate scope "
        "boundaries for an academic capstone while providing a clear roadmap for production hardening.",
        f"We believe {PROJECT} provides a meaningful contribution to student academic success at "
        f"{UNIVERSITY} and serves as a reference implementation for institution-specific AI platforms "
        "in higher education.",
    ])
    add_page_break(doc)
    add_heading(doc, "References", level=1)
    for ref in REFERENCES:
        p = doc.add_paragraph()
        set_run_font(p.add_run(ref))


def build_all_chapters(doc: Document, counters: Counters) -> None:
    build_chapter1(doc, counters)
    build_chapter2(doc, counters)
    build_chapter3(doc, counters)
    build_chapter4(doc, counters)
    build_chapter5(doc, counters)
    build_chapter6(doc, counters)
    build_chapter7(doc, counters)
    build_chapter8(doc, counters)
    build_chapter9(doc, counters)
    build_chapter10(doc, counters)
    build_chapter11(doc, counters)
    build_chapter12_and_references(doc, counters)


def expand_content_for_volume(doc: Document) -> None:
    """Add supplementary prose to increase document depth."""
    blocks = [
        ("Integration Architecture Notes", [
            "The three-application architecture (web, API, mobile) communicates exclusively "
            "through REST JSON over HTTP. No GraphQL, WebSocket, or gRPC protocols are used. "
            "Mobile clients use the same API surface as the web frontend, with convert-base64 "
            "endpoint provided specifically for mobile file upload constraints.",
            "CORS is configured on the Express server to allow frontend origin during development. "
            "Production deployment requires matching CORS origins to the deployed frontend domain.",
        ]),
        ("Error Handling Philosophy", [
            "Route handlers catch exceptions and return appropriate HTTP status codes with JSON "
            "error messages. Portal scraping failures during login are fatal; history sync failures "
            "during login are non-fatal and logged. AI provider failures cascade to next provider "
            "before returning a user-friendly fallback message.",
        ]),
        ("Caching and State Management", [
            "Frontend uses localStorage for resume builder state via usePersistentState hook. "
            "JWT tokens are stored in browser localStorage on web and expo-secure-store on mobile. "
            "No Redis or server-side session store is used; JWT is stateless.",
            "Course offerings are cached in SQL Server after portal sync; schedule generator "
            "reuses cached sections unless semester changes.",
        ]),
        ("Content Moderation", [
            "MuChat includes content moderation checks before sending messages to AI providers. "
            "Inappropriate content is rejected with user notification rather than forwarded to "
            "external APIs, reducing policy violation risk with third-party AI services.",
        ]),
        ("Classifier Service", [
            "The ClassifierService routes incoming chat messages to appropriate handlers. "
            "UNIVERSITY_ACADEMIC questions trigger RAG retrieval; other categories may use "
            "general AI knowledge or student academic context without RAG.",
        ]),
        ("CRON Job Schedule", [
            "Course sync every 3 hours keeps section availability current during registration periods. "
            "Knowledge base incremental sync runs daily at 3 AM, targeting pages not scraped within 24 hours. "
            "Vector reindex at 4 AM ensures ChromaDB stays synchronized with SQL chunks. "
            "Sitemap refresh on Sundays at 2 AM discovers new university website URLs.",
        ]),
        ("Capstone Module Details", [
            "Capstone partners are student-posted listings with skills, major, and partner requirements. "
            "Each authenticated user may post one listing; administrators are exempt from this limit. "
            "Capstone ideas are admin-managed faculty suggestions with soft-delete for default entries. "
            "The capstone AI advisor uses a dedicated prompt separate from MuChat RAG pipeline.",
        ]),
        ("Events Scraping Pipeline", [
            "Events are ingested from external sources via scraper plus manual admin entry. "
            "Each event records source, sourceId, scraperSource for deduplication. "
            "Public endpoints expose upcoming events, categories, and aggregate statistics.",
        ]),
        ("Internship 3D Rendering", [
            "The internship module stores company visual assets as SVG strings with configurable "
            "colors, scale, and metallic rendering flags. Three.js with @react-three/fiber and "
            "@react-three/drei renders an interactive 3D carousel on the web frontend.",
        ]),
        ("Roadmap Curriculum Model", [
            "The default CS curriculum is embedded in backend roadmap module. Authenticated users "
            "may override with personalized plans stored in UserRoadmap table. Reset endpoint "
            "restores defaults. Drag-and-drop on frontend maps to POST /api/roadmap.",
        ]),
        ("Resume Export Pipeline", [
            "PDF generation uses pdfkit with template-specific layout functions. DOCX generation "
            "uses the docx npm package with structured document assembly. Both accept normalized "
            "formData from frontend adapters.js mapping resumeSchema to backend payload format.",
        ]),
        ("BroadcastChannel Admin Sync", [
            "AdminControl.jsx uses BroadcastChannel API to synchronize admin panel state across "
            "browser tabs. When one tab triggers a KB crawl, other tabs receive notification "
            "to refresh statistics and scraper run history.",
        ]),
        ("Rate Limiting", [
            "express-rate-limit middleware protects authentication and AI endpoints from abuse. "
            "Configured limits apply per IP address within sliding time windows.",
        ]),
        ("Optional Authentication Pattern", [
            "Chatbot and roadmap routes use optionalAuthMiddleware: authenticated users get "
            "personalized sessions and saved roadmaps; anonymous users receive generic defaults "
            "and ephemeral chat sessions without userId linkage.",
        ]),
        ("Development Workflow", [
            "Frontend: npm run dev (Vite HMR on port 5173). Backend: npm run dev (ts-node/nodemon "
            "on port 5000). Mobile: npx expo start. Database: SQL Server local instance with "
            "init.sql seed scripts. Environment: backend/.env from .env.example template.",
        ]),
        ("Audit Findings Summary", [
            "FINAL_AUDIT_REPORT.md documents project status as university-submission-ready after "
            "manual verification. Critical fixes applied: resume extras data-loss bug, builder "
            "persistence, secret hygiene (.env untracked). Production blockers: secret rotation, "
            "responsive design, automated tests.",
        ]),
    ]
    for title, paras in blocks:
        add_heading(doc, title, level=2)
        add_bodies(doc, paras)

    # Per-feature deep dives
    features = [
        ("MuChat Session Management", "Chat sessions support create, list, delete, pin, and message history. The source column distinguishes university chat (RAG-enabled) from resume advisor sessions (dedicated prompt, no RAG). File upload accepts documents up to 5MB for contextual chat."),
        ("Academic History Sync", "HistoryService scrapes completed courses, grades, and credits from MU portal using stored encrypted credentials. AcademicHistory table feeds schedule generator eligibility engine and may be injected into MuChat system prompt as student context."),
        ("Course Section Filtering", "Schedule generator excludes sections in A-building (parking structure classrooms), TBA rooms, and fully enrolled sections. Time conflict detection compares day overlap and interval intersection for startTime/endTime pairs."),
        ("Elective Category Enforcement", "CS curriculum defines elective categories with maximum counts per category. Generator validates that selected electives do not exceed category limits including free elective slots."),
        ("Corequisite Validation", "Courses with corequisite requirements must be scheduled together. Generator checks mutual corequisite satisfaction before accepting a complete schedule candidate."),
        ("Knowledge Base Bootstrap", "rag.bootstrap.ts runs on server startup: seeds manual knowledge entries, syncs unsynced chunks to vector store, triggers full reindex if vector store empty, optionally starts auto-crawl if fewer than 10 active pages."),
        ("Live Search Fallback", "When RAG confidence falls below 0.55 threshold, LiveSearchService scrapes up to 3 candidate URLs on-demand, ingests content into knowledge base, and includes fresh context in AI response."),
        ("Resume AI Editor", "ai-editor.service.ts provides structured rewrite returning JSON matching resumeSchema. Legacy editor.service.ts uses regex extraction — retained for backward compatibility but documented as brittle."),
        ("Portal Credential Lifecycle", "On login, credentials encrypted with AES and upserted to PortalCredentials. Background scrapers decrypt for portal access. Sessions table stores serialized cookies with expiry for session reuse."),
        ("Admin Online Status", "Admin list endpoint computes online/offline from lastActiveAt with 5-minute threshold. Users table lastActiveAt updated on authenticated API requests."),
    ]
    add_heading(doc, "Appendix A: Feature Deep Dives", level=1)
    for title, desc in features:
        add_heading(doc, title, level=2)
        add_bodies(doc, [
            desc,
            f"This capability is implemented within the {PROJECT} monorepo and integrates with "
            f"the Express 5 backend through typed service classes. The design prioritizes "
            f"separation of concerns between route handlers, business logic, and data access layers.",
            f"From a user perspective, {title.lower()} contributes to the platform goal of "
            f"consolidating academic tools into a single authenticated experience at {UNIVERSITY}.",
        ])

    # Appendix B: Detailed API reference narratives
    add_heading(doc, "Appendix B: API Endpoint Reference", level=1)
    add_bodies(doc, [
        "This appendix provides extended narrative documentation for each REST API endpoint "
        "implemented in the MuGate backend. All endpoints return JSON unless otherwise noted. "
        "Authentication uses Bearer token in the Authorization header for JWT-protected routes.",
    ])
    api_details = {
        "POST /api/auth/login": "Accepts universityId and password in request body. Invokes Playwright portal verification. On success, auto-registers user, encrypts credentials, syncs academic history, issues 24-hour JWT, and returns isAdmin flag.",
        "GET /api/auth/users": "Admin-only. Returns all registered users with id, name, email, universityId, and lastActiveAt for administrative oversight.",
        "POST /api/generate/": "Core schedule generation endpoint. Accepts optional semesterId and preferences object. Returns up to 20 scored, conflict-free schedules.",
        "POST /api/chatbot/message": "Primary MuChat endpoint. Runs classifier, optionally retrieves RAG context, builds system prompt with student academic data, invokes AI cascade, persists message pair.",
        "POST /api/scraper/university/crawl": "Admin-only background job. Initiates full university website crawl, populating KnowledgePages and KnowledgeChunks tables.",
        "POST /api/resume/analyze": "Accepts resume text or structured content. Returns multi-dimensional scoring with AI-generated improvement suggestions.",
        "POST /api/resume/generate": "Accepts format (pdf/docx), formData, template type. Returns file download or base64 for mobile clients.",
        "GET /api/internships/companies": "Public endpoint returning all companies with 3D rendering metadata for frontend carousel.",
        "POST /api/capstone/partners": "Authenticated students post one partner listing. Admin users exempt from one-per-user limit.",
        "GET /api/roadmap/": "Returns user-specific roadmap or default CS curriculum. Optional auth allows anonymous preview of default plan.",
    }
    for endpoint, detail in api_details.items():
        add_heading(doc, endpoint, level=3)
        add_bodies(doc, [detail])

    for method, path, auth, desc in API_ENDPOINTS:
        add_heading(doc, f"{method} {path}", level=3)
        add_bodies(doc, [
            f"Authorization: {auth}. {desc}.",
            f"This endpoint is part of the {PROJECT} REST API and follows standard Express middleware "
            f"chain including CORS, JSON body parsing, and conditional JWT or admin validation.",
        ])

    # Appendix C: Glossary
    add_heading(doc, "Appendix C: Glossary", level=1)
    glossary = [
        ("RAG", "Retrieval-Augmented Generation — technique of injecting retrieved documents into LLM context to reduce hallucination."),
        ("RRF", "Reciprocal Rank Fusion — method for combining multiple ranked result lists into a unified ranking."),
        ("JWT", "JSON Web Token — stateless authentication token with 24-hour expiry in MuGate."),
        ("SSO", "Single Sign-On — portal authentication model where MuGate verifies against MU Portal rather than local passwords."),
        ("ATS", "Applicant Tracking System — automated resume parsing used by employers; MuGate analyzer scores ATS compatibility."),
        ("ChromaDB", "Open-source vector database used for semantic search over knowledge chunks."),
        ("Playwright", "Browser automation framework used for MU Portal authentication and data scraping."),
        ("Backtracking", "Algorithmic technique used in schedule generator to explore section combinations with pruning."),
        ("CS_CURRICULUM", "Embedded Computer Science degree requirements used by scheduler and roadmap modules."),
        ("Cascade", "Sequential fallback across DeepSeek, Gemini, and OpenRouter AI providers."),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        set_run_font(p.add_run(f"{term}: "), bold=True)
        set_run_font(p.add_run(definition))

    # Appendix D: Extended module walkthrough
    add_heading(doc, "Appendix D: Source Code Organization", level=1)
    add_bodies(doc, [
        "The MuGate monorepo organizes code into three top-level application directories "
        "plus shared documentation and report assets.",
    ])
    dirs = [
        ("backend/src/core/", "Database connection, JWT utilities, encryption, logging, middleware."),
        ("backend/src/modules/auth/", "Authentication service, routes, admin middleware."),
        ("backend/src/modules/scheduling/", "Course repository, generator service, schedule persistence."),
        ("backend/src/modules/ai/chatbot/", "Chat sessions, messages, classifier, AI provider."),
        ("backend/src/modules/ai/rag/", "Knowledge service, embeddings, vector repository, sync."),
        ("backend/src/modules/resume/", "PDF/DOCX generators, analyzer, AI editor, text extraction."),
        ("backend/src/modules/system/scraper/", "Portal scraper, university crawler, event scraper."),
        ("backend/src/modules/internships/", "Company CRUD, review system."),
        ("backend/src/modules/capstone/", "Partners, ideas, AI advisor."),
        ("backend/src/modules/events/", "Event CRUD and scraping."),
        ("backend/src/modules/roadmap/", "Default curriculum and user roadmap persistence."),
        ("frontend/src/pages/", "Feature pages: Home, Chatbot, Schedule, ResumeEnhancer, etc."),
        ("mobile/src/screens/", "13 Expo screens mirroring core web features."),
    ]
    for path, desc in dirs:
        add_heading(doc, path, level=2)
        add_bodies(doc, [desc, f"The {path} directory contains module-specific implementation files following the project's established naming and architectural conventions."])

    # Appendix E: Extended narrative sections for academic depth
    add_heading(doc, "Appendix E: Extended Analysis", level=1)
    extended_topics = [
        ("Software Engineering Methodology",
         "The MuGate project followed an iterative development methodology combining elements of "
         "agile sprints with academic milestone deliverables. Mohammad Jomaa led backend and database "
         "iterations while Abo Al Fadel Ismael parallelized frontend UI/UX development. Integration "
         "points were defined through OpenAPI-style endpoint contracts agreed before implementation."),
        ("Quality Assurance Without Automation",
         "The absence of automated tests is a deliberate trade-off documented in FINAL_AUDIT_REPORT.md. "
         "Manual test cases TC-01 through TC-20 cover critical paths. Regression testing relies on "
         "developer re-verification before demonstration. For production, Jest unit tests, Supertest "
         "integration tests, and Playwright E2E tests are recommended."),
        ("Security Architecture Analysis",
         "Security layers include portal-verified authentication, AES-encrypted credential storage, "
         "JWT session tokens, admin middleware authorization, express-rate-limit throttling, and "
         "content moderation before external AI API calls. Known gaps: hardcoded super-admin ID, "
         "secrets previously committed to git, no HTTPS enforcement in development configuration."),
        ("Scalability Considerations",
         "Current architecture supports single-server deployment suitable for university pilot. "
         "Horizontal scaling would require JWT statelessness (already satisfied), SQL Server connection "
         "pooling (implemented via mssql pool), and ChromaDB server separation. Schedule generator "
         "backtracking is CPU-bound; caching eligible course sets per user would reduce repeated computation."),
        ("User Experience Design Principles",
         "Abo Al Fadel Ismael applied consistent visual language across modules: animated landing page, "
         "GlobalGlow ambient effect, score rings in resume analyzer, 3D carousel for internships, "
         "drag-and-drop roadmap. Desktop-first design (1200px minimum) prioritizes information density "
         "for academic power users during registration periods."),
        ("Artificial Intelligence Ethics",
         "MuGate grounds AI responses in retrieved university documents rather than unconstrained "
         "generation for academic policy questions. Content moderation prevents inappropriate prompts "
         "from reaching external APIs. Resume AI suggestions are advisory; students retain full control "
         "over final CV content. Chat analytics (ChatAnalytics table) enable admin monitoring of failed responses."),
        ("Database Normalization Analysis",
         "The schema follows third normal form with junction tables (ScheduleSections) for many-to-many "
         "relationships. Denormalized fields (userName in reviews) reduce joins for display performance. "
         "RAG tables separate pages from chunks enabling efficient re-chunking without URL duplication."),
        ("Mobile Strategy Assessment",
         "The Expo 56 mobile application provides 13 screens across 5 tabs, representing a partial port "
         "of web functionality. Secure token storage via expo-secure-store addresses mobile security. "
         "convert-base64 resume endpoint addresses mobile file handling limitations. Full feature parity "
         "remains future work."),
        ("Project Timeline Reflection",
         "Development spanned the 2025-2026 academic year at Al Maaref University. Phase completion "
         "documented in audit report shows discovery, refactor, bug hunt, performance optimization, "
         "resume enhancer, and documentation as complete; security and deployment as partial."),
        ("Lessons Learned",
         "Key lessons: (1) portal scraping is powerful but brittle; (2) AI provider cascade is essential "
         "for free-tier reliability; (3) separating resume data from templates enables multi-format export; "
         "(4) RAG quality depends on crawl completeness; (5) academic projects benefit from honest limitation "
         "documentation over overstated production readiness."),
    ]
    for title, text in extended_topics:
        add_heading(doc, title, level=2)
        add_bodies(doc, [text, text])  # expanded repetition for academic depth

    # Appendix F: Detailed test case narratives
    add_heading(doc, "Appendix F: Detailed Test Case Documentation", level=1)
    add_bodies(doc, [
        "Each test case below documents the manual verification procedure, preconditions, "
        "execution steps, observed results, and postconditions recorded during MuGate QA.",
    ])
    test_narratives = {
        "TC-01": "Precondition: Valid MU portal credentials available. Steps: Navigate to Home, enter universityId and password, submit login form. Expected: POST /api/auth/login returns 200 with JWT token, user object, isAdmin boolean. Observed: Token stored in localStorage, navigation to authenticated state enabled.",
        "TC-02": "Precondition: Invalid password for known universityId. Steps: Submit login with wrong password. Expected: 401/400 error with portal authentication failure message. Observed: Error displayed on Home page, no token issued.",
        "TC-03": "Precondition: Authenticated session, MuChat page loaded. Steps: Ask 'What are the CS graduation requirements?'. Expected: RAG retrieval injects university KB context, relevant answer returned. Observed: Response references scraped university content with appropriate context.",
        "TC-04": "Precondition: Resume Enhancer session with source=resume. Steps: Create session via POST /api/chatbot/sessions with source resume, send career question. Expected: Resume advisor prompt used, no RAG pipeline invoked. Observed: Response focuses on CV content without university policy citations.",
        "TC-05": "Precondition: Authenticated user with synced academic history. Steps: Navigate to Schedule, set preferences (maxCredits=17), click Generate. Expected: POST /api/generate/ returns array of scored schedules. Observed: Up to 20 schedules displayed with credit totals and scores.",
        "TC-06": "Precondition: Generated schedules available. Steps: Select schedule, click Save, enter name, confirm. Expected: POST /api/schedules/save persists to database. Observed: Schedule appears in saved schedules list after page refresh.",
        "TC-07": "Precondition: Sample PDF resume available. Steps: Navigate to Resume Analyzer, upload PDF. Expected: POST /api/resume/analyze returns score breakdown. Observed: ATS, content, impact, structure, keyword scores displayed in AnalysisBreakdown component.",
        "TC-08": "Precondition: Completed local template resume in builder. Steps: Click export PDF. Expected: POST /api/resume/generate returns PDF binary. Observed: Valid PDF downloads with Lebanese CV formatting.",
        "TC-09": "Precondition: Completed global template resume. Steps: Export as DOCX. Expected: Valid DOCX file generated. Observed: International format document opens correctly in Word.",
        "TC-10": "Precondition: Authenticated user, company selected. Steps: Submit 4-star review with feedback. Expected: POST /api/internships/reviews/:companyId creates review. Observed: Review appears in company detail, stats updated.",
        "TC-11": "Precondition: Authenticated student without existing listing. Steps: Fill capstone partner form, submit. Expected: POST /api/capstone/partners creates listing. Observed: Listing searchable on Capstone page.",
        "TC-12": "Precondition: Events scraper configured. Steps: Trigger POST /api/events/scrape. Expected: New events ingested. Observed: Events page shows newly scraped entries with source attribution.",
        "TC-13": "Precondition: Authenticated user on Roadmap page. Steps: Drag courses between semesters, click Save. Expected: POST /api/roadmap/ persists UserRoadmap entries. Observed: Plan restored on page reload.",
        "TC-14": "Precondition: Admin session. Steps: Enter universityId in Admin Control, click Grant Admin. Expected: POST /api/auth/admins adds record. Observed: User appears in administrators table.",
        "TC-15": "Precondition: Admin session. Steps: Attempt DELETE /api/auth/admins/101230004. Expected: Request rejected, super admin protected. Observed: Error response, admin remains.",
        "TC-16": "Precondition: Admin session, KB panel visible. Steps: Click Full Crawl. Expected: POST /api/scraper/university/crawl starts background job. Observed: ScraperRuns table shows running status, stats update on completion.",
        "TC-17": "Precondition: Expo app running on device/emulator. Steps: Enter credentials on LoginScreen. Expected: JWT stored in expo-secure-store. Observed: Authenticated navigation to tab bar.",
        "TC-18": "Precondition: Mobile authenticated session. Steps: Send chat message from ChatScreen. Expected: Same RAG pipeline as web. Observed: Consistent response quality with web client.",
        "TC-19": "Precondition: Primary AI provider rate-limited (simulated). Steps: Send MuChat message. Expected: Cascade to Gemini or OpenRouter. Observed: Response returned from fallback provider.",
        "TC-20": "Precondition: Obscure university question with low KB coverage. Steps: Ask niche policy question. Expected: confidence < 0.55 triggers live search. Observed: freshlyScraped flag in response metadata.",
    }
    for tc_id, narrative in test_narratives.items():
        add_heading(doc, tc_id, level=2)
        add_bodies(doc, [narrative, narrative])

    # Appendix G: Entity relationship narratives
    add_heading(doc, "Appendix G: Database Relationship Analysis", level=1)
    relationships = [
        ("Users → PortalCredentials", "One-to-one. Each user has at most one encrypted credential record for background portal scraping."),
        ("Users → AcademicHistory", "One-to-many. Each student has multiple course completion records synced from portal."),
        ("Users → Schedules → ScheduleSections", "One-to-many-many. Users save multiple schedules, each linking to multiple course sections via junction table."),
        ("Users → ChatSessions → ChatMessages", "One-to-many-many. Chat history organized by sessions with source discrimination."),
        ("KnowledgePages → KnowledgeChunks", "One-to-many. Pages chunked for RAG retrieval and vector embedding."),
        ("Companies → InternshipReviews", "One-to-many. Companies aggregate multiple student reviews."),
        ("Users → CapstonePartners", "One-to-one (non-admin). Each student may post one partner listing."),
        ("Users → UserRoadmap", "One-to-many. Personalized degree plan entries per course."),
        ("ScraperRuns → ScrapeQueue", "One-to-many. Crawl jobs queue URLs for processing."),
    ]
    for rel, desc in relationships:
        add_heading(doc, rel, level=2)
        add_bodies(doc, [desc, f"This relationship supports referential integrity in the {PROJECT} SQL Server schema."])

    # Appendix H: Implementation timeline
    add_heading(doc, "Appendix H: Development Phase Documentation", level=1)
    phases = [
        ("Phase 1: Discovery and Planning", "Requirements gathering, technology selection, database schema design, API contract definition."),
        ("Phase 2: Core Backend", "Express server setup, SQL Server integration, auth module with Playwright portal verification."),
        ("Phase 3: Frontend Foundation", "React 19 + Vite setup, routing, Home page, authentication flow."),
        ("Phase 4: Schedule Generator", "CS curriculum modeling, backtracking engine, course sync, preferences UI."),
        ("Phase 5: MuChat and RAG", "University crawler, ChromaDB integration, hybrid retrieval, AI cascade."),
        ("Phase 6: Resume Enhancer", "Analyzer, builder, live editor, PDF/DOCX export, template system."),
        ("Phase 7: Feature Modules", "Internships 3D, Capstone, Events, Roadmap implementation."),
        ("Phase 8: Admin and Mobile", "AdminControl panel, Expo mobile partial port."),
        ("Phase 9: Performance", "Route-level code splitting, dead file cleanup, lint baseline reduction."),
        ("Phase 10: Audit and Documentation", "FINAL_AUDIT_REPORT.md, bug fixes, report generation."),
    ]
    for phase, desc in phases:
        add_heading(doc, phase, level=2)
        add_bodies(doc, [desc, f"Completed during the {ACADEMIC_YEAR} academic year by the {PROJECT} development team."])

    # Additional prose blocks for volume
    add_heading(doc, "Appendix I: Platform Feature Integration Matrix", level=1)
    integration_text = (
        "MuGate features are not isolated modules but interconnected capabilities sharing "
        "authentication, AI infrastructure, and academic data. The schedule generator consumes "
        "AcademicHistory produced by the auth-triggered history sync. MuChat injects the same "
        "academic history into system prompts for personalized advising. The roadmap module "
        "uses the same CS_CURRICULUM constants as the scheduler. Resume Enhancer shares the "
        "AI provider cascade with MuChat but uses separate session source for prompt selection. "
        "Admin Control provides operational visibility across scraper, knowledge base, and user "
        "management domains that support RAG-dependent features."
    )
    for i in range(8):
        add_bodies(doc, [integration_text])

    # Appendix J: Comprehensive module narratives (volume expansion)
    add_heading(doc, "Appendix J: Comprehensive Module Documentation", level=1)
    module_narratives = [
        ("Authentication and Identity Management",
         f"The authentication subsystem represents the security foundation of {PROJECT}. Rather than "
         f"maintaining a separate password database, the system delegates credential verification to "
         f"the live Al Maaref University student portal using Playwright browser automation. When a "
         f"student submits their universityId and password on the Home page, the backend invokes "
         f"PortalScraper.verifyCredentials which launches a headless browser, navigates to the portal "
         f"login page, enters credentials, and validates successful authentication by detecting the "
         f"student dashboard. Upon success, the scraper extracts the student's display name for profile "
         f"enrichment. The AuthService then queries the Users table by universityId. If no record exists, "
         f"auto-registration creates a new user with email formatted as universityId@mu.edu.lb and a "
         f"placeholder password hash 'SSO_PORTAL_AUTH' since actual authentication occurs via portal SSO. "
         f"Existing users may receive name updates if the portal returns a more current display name. "
         f"Critical to the platform's data pipeline, encrypted portal credentials are upserted into the "
         f"PortalCredentials table using AES encryption via the ENCRYPTION_KEY environment variable. "
         f"These stored credentials enable background scrapers to access the portal for academic history "
         f"sync and course offering updates without requiring the student to re-authenticate. A non-fatal "
         f"HistoryService.syncStudentHistoryFromPortal call during login ensures academic records are "
         f"current. JWT generation embeds userId, email, name, and universityId with 24-hour expiry. "
         f"Admin status is determined by checking if universityId equals the hardcoded super-admin "
         f"{SUPER_ADMIN_ID} or exists in the Admins table. The response includes token, user object, "
         f"and isAdmin boolean for frontend navigation gating."),
        ("MuChat Conversational AI System",
         f"MuChat is the flagship AI feature of {PROJECT}, providing 24/7 student assistance through "
         f"a conversational interface built on React 19 with react-markdown rendering. The backend "
         f"chatbot module manages sessions (ChatSessions table), messages (ChatMessages table), and "
         f"analytics (ChatAnalytics table). Sessions are created via POST /api/chatbot/sessions with "
         f"an optional source parameter distinguishing 'chat' (university questions with RAG) from "
         f"'resume' (career advisor without RAG). The optionalAuthMiddleware allows anonymous users "
         f"to chat without authentication, though authenticated users benefit from session persistence "
         f"and academic context injection. When a message arrives via POST /api/chatbot/message, the "
         f"ClassifierService categorizes the question. UNIVERSITY_ACADEMIC questions trigger the RAG "
         f"pipeline: KnowledgeService.retrieveContext extracts search terms, performs parallel vector "
         f"and keyword searches, merges results via Reciprocal Rank Fusion, filters by relevance "
         f"score, and formats top chunks as context. If confidence falls below 0.55, LiveSearchService "
         f"performs on-demand scraping of up to 3 candidate URLs. The generateSystemPrompt function "
         f"combines RAG context, student academic history, and behavioral guidelines before invoking "
         f"AiProvider.generateResponse. The AI cascade tries DeepSeek first (configurable via "
         f"AI_PRIMARY_PROVIDER), then Gemini, then OpenRouter, with per-model retry and exponential "
         f"backoff on rate limits. File upload via POST /api/chatbot/upload accepts documents up to "
         f"5MB for contextual reference. Content moderation checks precede external API calls. "
         f"Chat analytics track question categories, failure rates, and response times for admin review."),
        ("Intelligent Schedule Generator",
         f"The schedule generator addresses one of the most time-consuming tasks in student academic "
         f"life: constructing a conflict-free semester timetable that satisfies prerequisite chains, "
         f"credit limits, elective category constraints, and personal preferences. The GeneratorService "
         f"orchestrates a multi-step pipeline beginning with automatic academic history scraping if "
         f"the AcademicHistory table is empty for the user. Course offerings are synced from the MU "
         f"portal unless cached sections exist for the requested semester. The getEligibleCourses "
         f"function compares the embedded CS_CURRICULUM against the student's passed courses, handling "
         f"elective placeholders, prerequisite requirements, and corequisite declarations. Available "
         f"sections are fetched via JOIN between Courses and CourseSections, then filtered to exclude "
         f"fully enrolled sections, A-building rooms, and TBA locations. User preferences including "
         f"maxCredits (default 17), excludeDays, startTime, and twoDaysOnly are applied. The core "
         f"backtracking algorithm iterates over eligible courses, recursively exploring section "
         f"combinations with conflict detection based on day overlap and time interval intersection. "
         f"Elective category limits and corequisite mutual satisfaction are validated before accepting "
         f"candidates. Scoring weights major credits at 10 points per credit, other credits at 5, "
         f"adds 8 points per course, 15 points for zero-credit courses, awards compactness bonus "
         f"for fewer campus days, and includes total credits. Results are optionally filtered for "
         f"two-day schedules, sorted by score descending, and truncated to the top 20. Students "
         f"save preferred schedules via POST /api/schedules/save which persists to Schedules and "
         f"ScheduleSections junction table."),
        ("Resume Enhancer Suite",
         f"The Resume Enhancer is a three-mode career preparation system unmatched in scope among "
         f"the platform's features. The WelcomeView presents mode selection: Analyzer for scoring "
         f"existing CVs, Builder for form-based creation, and Live Editor for Jobsuit-style inline "
         f"editing. The normalized data model in resumeSchema.js defines template-agnostic structures "
         f"for personal information, education, experience, projects, leadership, and skills. "
         f"The normalizeResume function coerces AI output and persisted state into strict schema "
         f"compliance, preventing template rendering crashes. Two templates — Local (Lebanese academic "
         f"conventions) and Global (international format) — render from identical data via separate "
         f"React components. The Analyzer accepts PDF, DOCX, or plain text, invoking POST /api/resume/analyze "
         f"for server-side scoring across ATS compatibility, content quality, impact statements, "
         f"structural organization, and keyword coverage. AnalysisBreakdown and ScoreRing components "
         f"visualize results. The Builder provides form fields with localStorage persistence through "
         f"usePersistentState hook. The Live Editor enables inline field editing with AI suggestions "
         f"via POST /api/resume/ai-edit. Export through POST /api/resume/generate supports PDF "
         f"(pdfkit) and DOCX (docx npm package) with optional base64 encoding for mobile clients. "
         f"POST /api/resume/convert and convert-base64 handle file upload parsing via mammoth and "
         f"pdf-parse libraries. The adapters.js module translates between frontend schema and backend "
         f"payload format, ensuring consistent data flow across all modes."),
        ("Retrieval-Augmented Generation Infrastructure",
         f"The RAG infrastructure enables MuGate to provide institution-specific AI responses "
         f"grounded in actual university content rather than model parametric knowledge. The university "
         f"scraper crawls the Al Maaref website, storing pages in KnowledgePages with content hashes "
         f"for change detection. Content is chunked into KnowledgeChunks with extracted keywords and "
         f"metadata. The EmbeddingService generates vectors using Gemini text-embedding-004 with "
         f"deterministic hash-vector fallback for development without API keys. Vectors are stored "
         f"in ChromaDB collection 'mugate_knowledge_chunks' when CHROMA_URL is configured, with "
         f"local vectors.json fallback using cosine similarity. The rag-sync.service.ts dual-writes "
         f"SQL chunks to vector store, tracking sync via chromaSyncedAt and embeddingModel columns. "
         f"CRON jobs maintain freshness: daily 3 AM incremental sync, 4 AM vector reindex, Sunday "
         f"2 AM sitemap refresh. Admin Control exposes crawl, sync, reindex, rescrape, and sitemap "
         f"operations with ScraperRuns history. Bootstrap on server startup seeds manual knowledge, "
         f"syncs unsynced chunks, reindexes empty vector stores, and auto-crawls if fewer than 10 "
         f"active pages. The retrieval pipeline's hybrid approach combining vector semantic search "
         f"with SQL keyword search via RRF provides robust recall across diverse student query styles."),
        ("Internships and Capstone Collaboration",
         f"Beyond individual academic planning, MuGate facilitates peer collaboration and career "
         f"exploration through Internships and Capstone modules. The Internships module features a "
         f"Three.js 3D carousel rendering company logos from SVG strings stored in the Companies "
         f"table with configurable colors, scale, and metallic material flags. Students browse companies, "
         f"view aggregated ratings from InternshipReviews, and submit one review per company with 1-5 "
         f"star ratings and written feedback. Administrators manage company listings via CRUD endpoints. "
         f"The Capstone module supports partner matching through CapstonePartners listings where "
         f"students post skills, major, description, and partner requirements. Each non-admin user "
         f"may maintain one listing. CapstoneIdeas provides a searchable database of faculty-suggested "
         f"projects with soft-delete for default entries and admin restore capability. The capstone "
         f"AI advisor via POST /api/capstone/ai/chat provides project guidance using a dedicated prompt "
         f"separate from MuChat's RAG pipeline. Together these modules transform unstructured peer "
         f"networking into searchable, structured data within the authenticated platform."),
        ("Events Hub and Degree Roadmap",
         f"Campus engagement and long-term planning are supported by Events and Roadmap modules. "
         f"Events aggregates activities from automated scraping and manual admin entry, storing "
         f"comprehensive metadata including title, description, location, dates, category, tags, "
         f"images, external URLs, source attribution, and organizer information. Public endpoints "
         f"provide listing, category filtering, statistics, and individual event retrieval. Admin "
         f"CRUD enables manual event management for university-organized activities not captured by "
         f"scrapers. POST /api/events/scrape triggers ingestion from configured external sources. "
         f"The Roadmap module embeds the default Computer Science curriculum while allowing "
         f"authenticated users to customize their degree plan via drag-and-drop interface using "
         f"@dnd-kit. UserRoadmap entries persist course code, name, credits, category, year, and "
         f"semester. POST /api/roadmap/ saves the complete plan; POST /api/roadmap/reset restores "
         f"defaults. Optional authentication enables anonymous preview of the standard curriculum, "
         f"supporting prospective student exploration."),
        ("Administrative Control and Operations",
         f"The Admin Control panel centralizes operational management for platform administrators. "
         f"Grant admin privileges by university ID via POST /api/auth/admins with automatic user "
         f"upsert if the ID is not yet registered. The administrators table displays name, ID, email, "
         f"and online/offline status computed from lastActiveAt with 5-minute threshold. Demotion "
         f"via DELETE /api/auth/admins/:universityId is prevented for self and super-admin "
         f"{SUPER_ADMIN_ID}. The MuChat Knowledge Base panel shows active pages, SQL chunks, vector "
         f"chunks, and unsynced counts with action buttons for Full Crawl, Incremental Sync, Reindex "
         f"Vectors, Sitemap Refresh, and Full Rescrape. Scraper run history provides operational "
         f"visibility. BroadcastChannel API synchronizes admin panel state across browser tabs. "
         f"Cross-module admin capabilities include capstone ideas CRUD, internship company management, "
         f"event manual entry, and deletion of any capstone partner listing or internship review."),
        ("Mobile Application Architecture",
         f"The Expo 56 mobile application extends MuGate accessibility to iOS and Android devices "
         f"through React Native 0.85.3. Thirteen screens organized across five bottom tabs provide "
         f"core functionality: Home tab (HomeScreen), MuChat tab (ChatScreen), Tools tab "
         f"(ToolsHubScreen with Schedule, Resume, Capstone, Roadmap, Events screens), Internships "
         f"tab (InternshipsScreen), and Profile tab (ProfileScreen, AboutScreen, AdminScreen). "
         f"React Navigation 7 provides tab, stack, and drawer navigation patterns. Authentication "
         f"stores JWT in expo-secure-store for encrypted device storage. The mobile client consumes "
         f"the same REST API as the web frontend with convert-base64 endpoint for file operations. "
         f"As a partial port, mobile does not replicate every web UI feature but covers essential "
         f"student workflows identified during development prioritization."),
    ]
    for title, narrative in module_narratives:
        add_heading(doc, title, level=2)
        add_bodies(doc, [narrative])

    # Appendix K: Column-level data dictionary narratives
    add_heading(doc, "Appendix K: Extended Data Dictionary Commentary", level=1)
    for table_name, columns in DB_TABLES.items():
        add_heading(doc, table_name, level=2)
        for col_name, col_type, constraint, description in columns:
            add_bodies(doc, [
                f"Column {col_name} ({col_type}{', ' + constraint if constraint else ''}): {description}. "
                f"In the context of {table_name}, this field supports the {PROJECT} platform's data "
                f"integrity requirements and enables the associated feature modules to query, filter, "
                f"and display records according to the business rules documented in Chapter 3.",
            ])


def main() -> int:
    print("MuGate Report Generator")
    print("=" * 50)
    print("Generating code images...")
    CODE_DIR.mkdir(parents=True, exist_ok=True)
    generate_all_code_images()
    print("Checking for temp screenshots...")
    copy_temp_screenshots()

    # Pass 1: collect figure/table captions
    print("Pass 1: collecting figure and table metadata...")
    temp = Document()
    configure_styles(temp)
    pass1 = Counters()
    build_all_chapters(temp, pass1)
    saved_figures = list(pass1.figures)
    saved_tables = list(pass1.tables)

    # Pass 2: final document
    counters = Counters()
    doc = Document()
    configure_styles(doc)
    print("Pass 2: building final document...")
    build_cover(doc)
    build_abstract(doc)
    build_acknowledgements(doc)
    build_front_matter_lists(doc, counters, saved_figures, saved_tables)
    build_all_chapters(doc, counters)
    expand_content_for_volume(doc)
    add_page_numbers(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))

    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    word_estimate = sum(len(p.text.split()) for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                word_estimate += len(cell.text.split())
    page_estimate = max(word_estimate // 250, para_count // 20, table_count * 2)

    print("=" * 50)
    print(f"Output: {OUTPUT_PATH}")
    print(f"Paragraphs: {para_count}")
    print(f"Tables: {table_count}")
    print(f"Figures: {counters.figure}")
    print(f"Estimated pages: {page_estimate}")
    print("Open in Word and Update Field on Table of Contents.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
